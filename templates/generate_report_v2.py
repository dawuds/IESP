#!/usr/bin/env python3
"""
Generate IESP Report Template v2 (.docx)
Aligned with BNM RMiT Appendix 7 Part A format and v4 AWP structure.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colour constants ──
DARK_BLUE = RGBColor(0x1F, 0x3A, 0x5F)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GREY = RGBColor(0x33, 0x33, 0x33)

HEADER_BG_HEX = "1F3A5F"
LIGHT_GREY_HEX = "F2F2F2"
GREEN_HEX = "E2EFDA"
YELLOW_HEX = "FFF2CC"


def set_cell_shading(cell, colour_hex):
    """Set cell background colour."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colour_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, left, right, insideH, insideV."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)
    for edge, val in kwargs.items():
        element = tcBorders.find(qn(f'w:{edge}'))
        if element is None:
            element = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
                f'w:sz="{val.get("sz", "4")}" w:space="0" '
                f'w:color="{val.get("color", "000000")}"/>'
            )
            tcBorders.append(element)


def styled_table(doc, rows, cols, col_widths=None):
    """Create a table with professional styling."""
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def style_header_row(row, font_size=9):
    """Style a table header row with dark blue background and white text."""
    for cell in row.cells:
        set_cell_shading(cell, HEADER_BG_HEX)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(font_size)
                run.font.name = 'Calibri'


def style_data_row(row, font_size=9, alternate=False):
    """Style a data row."""
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        if alternate:
            set_cell_shading(cell, LIGHT_GREY_HEX)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(font_size)
                run.font.name = 'Calibri'
                run.font.color.rgb = DARK_GREY


def add_field_row(table, label, value="", row_idx=None):
    """Add a label-value row to an existing table."""
    if row_idx is not None:
        row = table.rows[row_idx]
    else:
        row = table.add_row()
    row.cells[0].text = label
    for p in row.cells[0].paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_GREY
    row.cells[1].text = value
    for p in row.cells[1].paragraphs:
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_GREY
    return row


def add_section_header_row(table, text):
    """Add a section header row spanning the full width of a 2-col table."""
    row = table.add_row()
    row.cells[0].merge(row.cells[1])
    row.cells[0].text = text
    set_cell_shading(row.cells[0], LIGHT_GREY_HEX)
    for p in row.cells[0].paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_BLUE


def add_heading(doc, text, level=1):
    """Add a heading with dark blue colour."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK_BLUE
        run.font.name = 'Calibri'
    return h


def add_para(doc, text, bold=False, italic=False, size=10, color=None, space_after=6, alignment=None):
    """Add a paragraph with styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color or DARK_GREY
    p.paragraph_format.space_after = Pt(space_after)
    if alignment:
        p.alignment = alignment
    return p


def add_placeholder(doc, text, size=10):
    """Add placeholder text in square brackets."""
    return add_para(doc, text, italic=True, size=size, color=RGBColor(0x66, 0x66, 0x66))


def set_section_margins(section):
    """Set professional margins."""
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_footer(section, text):
    """Add footer with text and page number."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(text + "  |  Page ")
    run.font.name = 'Calibri'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Add page number field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run2 = p.add_run()
    run2._r.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run3 = p.add_run()
    run3._r.append(instrText)
    run3.font.name = 'Calibri'
    run3.font.size = Pt(8)
    run3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run4 = p.add_run()
    run4._r.append(fldChar2)


def build_finding(doc, finding_id, title, domain, bnm_ref, level, platform, risk_rating, conclusion,
                  condition, criteria, cause, effect, recommendation,
                  agreed_action, responsible, target_date, status):
    """Build a single finding block."""
    add_heading(doc, f"Finding [{finding_id}]: {title}", level=3)

    # Attributes table
    table = styled_table(doc, rows=7, cols=2, col_widths=[2.0, 4.5])
    headers = ["Attribute", "Detail"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    style_header_row(table.rows[0])

    data = [
        ("Domain", domain),
        ("BNM Reference", bnm_ref),
        ("Assessment Level", level),
        ("Platform/System", platform),
        ("Risk Rating", risk_rating),
        ("Conclusion", conclusion),
    ]
    for idx, (attr, detail) in enumerate(data, 1):
        table.rows[idx].cells[0].text = attr
        table.rows[idx].cells[1].text = detail
        for p in table.rows[idx].cells[0].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Calibri'
        for p in table.rows[idx].cells[1].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = 'Calibri'
        if idx % 2 == 0:
            style_data_row(table.rows[idx], alternate=True)

    doc.add_paragraph()  # spacer

    # Condition/Criteria/Cause/Effect/Recommendation
    sections = [
        ("Condition:", condition),
        ("Criteria:", criteria),
        ("Cause:", cause),
        ("Effect:", effect),
        ("Recommendation:", recommendation),
    ]
    for label, text in sections:
        p = doc.add_paragraph()
        run_label = p.add_run(label + " ")
        run_label.font.bold = True
        run_label.font.name = 'Calibri'
        run_label.font.size = Pt(10)
        run_label.font.color.rgb = DARK_GREY
        run_text = p.add_run(text)
        run_text.font.name = 'Calibri'
        run_text.font.size = Pt(10)
        run_text.font.color.rgb = DARK_GREY
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()  # spacer

    # Management Response table
    p = doc.add_paragraph()
    run = p.add_run("Management Response:")
    run.font.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_BLUE

    table2 = styled_table(doc, rows=5, cols=2, col_widths=[2.0, 4.5])
    headers2 = ["Attribute", "Detail"]
    for i, h in enumerate(headers2):
        table2.rows[0].cells[i].text = h
    style_header_row(table2.rows[0])

    mgmt_data = [
        ("Agreed Action", agreed_action),
        ("Responsible Person", responsible),
        ("Target Date", target_date),
        ("Status", status),
    ]
    for idx, (attr, detail) in enumerate(mgmt_data, 1):
        table2.rows[idx].cells[0].text = attr
        table2.rows[idx].cells[1].text = detail
        for p in table2.rows[idx].cells[0].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Calibri'
        for p in table2.rows[idx].cells[1].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = 'Calibri'

    doc.add_paragraph()  # spacer


def main():
    doc = Document()

    # ── Set default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = DARK_GREY

    # Style headings
    for i in range(1, 5):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = DARK_BLUE

    # ── Section setup ──
    section = doc.sections[0]
    set_section_margins(section)
    add_footer(section, "Confidential — [Entity Name] — IESP Assessment Report")

    # ════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════════════════════════════

    # Spacer
    for _ in range(6):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Risk Assessment Report")
    run.font.name = 'Calibri'
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Per Appendix 7 Part A — BNM RMiT")
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("(BNM/RH/PD 028-98, 28 November 2025)")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Horizontal line
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 60)
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(10)

    doc.add_paragraph()

    # Cover page fields table
    cover_table = styled_table(doc, rows=1, cols=2, col_widths=[2.5, 4.0])
    # Remove header row styling for cover
    cover_data = [
        ("Entity Name", "[Name of Financial Institution]"),
        ("Engagement Type", "[Cloud / Emerging Tech / DCRA / NRA / Digital Services]"),
        ("Engagement Mode", "[Design Adequacy / Operating Effectiveness]"),
        ("Engagement Reference", "[IESP-YYYY-NNN]"),
        ("Report Date", "[DD/MM/YYYY]"),
        ("Report Version", "[Draft / Final]"),
        ("Classification", "CONFIDENTIAL / SULIT"),
    ]

    # Set first row
    cover_table.rows[0].cells[0].text = cover_data[0][0]
    cover_table.rows[0].cells[1].text = cover_data[0][1]
    for p in cover_table.rows[0].cells[0].paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_BLUE
    for p in cover_table.rows[0].cells[1].paragraphs:
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Calibri'

    # Add remaining rows
    for label, value in cover_data[1:]:
        row = cover_table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        for p in row.cells[0].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                run.font.color.rgb = DARK_BLUE
        for p in row.cells[1].paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'

    # Document control
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Document Control")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK_BLUE

    dc_table = styled_table(doc, rows=4, cols=4, col_widths=[1.0, 1.5, 2.0, 2.0])
    dc_headers = ["Version", "Date", "Author", "Changes"]
    for i, h in enumerate(dc_headers):
        dc_table.rows[0].cells[i].text = h
    style_header_row(dc_table.rows[0])

    dc_data = [
        ("0.1", "[DD/MM/YYYY]", "[Lead Assessor]", "Initial draft"),
        ("0.2", "[DD/MM/YYYY]", "[Peer Reviewer]", "Peer review comments incorporated"),
        ("1.0", "[DD/MM/YYYY]", "[Lead Assessor]", "Final version issued"),
    ]
    for idx, (ver, date, author, changes) in enumerate(dc_data, 1):
        dc_table.rows[idx].cells[0].text = ver
        dc_table.rows[idx].cells[1].text = date
        dc_table.rows[idx].cells[2].text = author
        dc_table.rows[idx].cells[3].text = changes
        style_data_row(dc_table.rows[idx], alternate=(idx % 2 == 0))

    # Page break
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ════════════════════════════════════════════════════════════════════

    add_heading(doc, "Table of Contents", level=1)

    toc_items = [
        "Section 1: Financial Institution",
        "Section 2: External Service Provider",
        "Section 3: Detail of Application",
        "Section 4: Technology Risk Assessment",
        "    4.1 Executive Summary",
        "    4.2 Assessment Summary by Control Domain",
        "    4.3 Detailed Findings",
        "    4.4 Compliance Summary",
        "    4.5 Findings by Risk Rating",
        "Section 5: Quality Assurance",
        "Section 6: Authorised Signatory",
        "",
        "Appendix I: Part C — Negative Attestation Statement",
        "Appendix II: Part B — Format of Confirmation",
        "Appendix III: Part D — Minimum Controls Summary",
        "Appendix IV: Document Request List",
    ]
    for item in toc_items:
        if item == "":
            doc.add_paragraph()
            continue
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        if item.startswith("    "):
            p.paragraph_format.left_indent = Inches(0.5)
            run.font.color.rgb = DARK_GREY
        elif item.startswith("Appendix"):
            run.font.color.rgb = DARK_BLUE
            run.font.bold = True
        else:
            run.font.color.rgb = DARK_BLUE
            run.font.bold = True
        p.paragraph_format.space_after = Pt(2)

    add_para(doc, "[Update page numbers before issuing final report]", italic=True, size=9,
             color=RGBColor(0x99, 0x99, 0x99))

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # SECTION 1: FINANCIAL INSTITUTION
    # ════════════════════════════════════════════════════════════════════

    add_heading(doc, "Section 1: Financial Institution", level=1)

    s1_table = styled_table(doc, rows=1, cols=2, col_widths=[2.8, 3.7])
    # Header
    s1_table.rows[0].cells[0].text = "Field"
    s1_table.rows[0].cells[1].text = "Details"
    style_header_row(s1_table.rows[0])

    s1_fields = [
        ("Name of Financial Institution", "[Full legal name]"),
        ("Mailing address", "[Registered address]"),
        ("Type of digital services / cloud service / emerging technology",
         "[New / Enhancement]"),
        ("Description of the digital services / cloud service / emerging technology",
         "[Provide a concise description of the service, platform, or technology "
         "being assessed. Include the business purpose and key capabilities.]"),
        ("Key contact personnel", "[Name, Title]"),
        ("Email address", "[email@institution.com.my]"),
        ("Phone number", "[+60-X-XXXX XXXX]"),
    ]
    for label, value in s1_fields:
        add_field_row(s1_table, label, value)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # SECTION 2: EXTERNAL SERVICE PROVIDER
    # ════════════════════════════════════════════════════════════════════

    add_heading(doc, "Section 2: External Service Provider", level=1)

    s2_table = styled_table(doc, rows=1, cols=2, col_widths=[2.8, 3.7])
    s2_table.rows[0].cells[0].text = "Field"
    s2_table.rows[0].cells[1].text = "Details"
    style_header_row(s2_table.rows[0])

    s2_fields = [
        ("Name of company", "[IESP firm legal name]"),
        ("SSM registration number", "[Company No.]"),
        ("Mailing address", "[Registered address]"),
        ("Engagement period", "[DD/MM/YYYY] to [DD/MM/YYYY]"),
        ("Key contact personnel", "[Name, Title]"),
        ("Email address", "[email@iesp.com]"),
        ("Phone number", "[+60-X-XXXX XXXX]"),
        ("Lead assessor name", "[Full name]"),
        ("Lead assessor qualifications",
         "[CISA, CISSP, ISO 27001 LA, or equivalent — list relevant certifications]"),
    ]
    for label, value in s2_fields:
        add_field_row(s2_table, label, value)

    doc.add_paragraph()

    # Team composition table
    add_heading(doc, "Assessment Team Composition", level=2)

    team_table = styled_table(doc, rows=5, cols=4, col_widths=[0.5, 2.0, 1.5, 2.5])
    team_headers = ["#", "Name", "Role", "Qualifications"]
    for i, h in enumerate(team_headers):
        team_table.rows[0].cells[i].text = h
    style_header_row(team_table.rows[0])

    team_data = [
        ("1", "[Name]", "Lead Assessor", "[CISA, CISSP, etc.]"),
        ("2", "[Name]", "Senior Assessor", "[Qualifications]"),
        ("3", "[Name]", "Assessor", "[Qualifications]"),
        ("4", "[Name]", "Subject Matter Expert", "[Qualifications]"),
    ]
    for idx, (num, name, role, quals) in enumerate(team_data, 1):
        team_table.rows[idx].cells[0].text = num
        team_table.rows[idx].cells[1].text = name
        team_table.rows[idx].cells[2].text = role
        team_table.rows[idx].cells[3].text = quals
        style_data_row(team_table.rows[idx], alternate=(idx % 2 == 0))

    doc.add_paragraph()

    # Independence declaration
    add_heading(doc, "Independence Declaration", level=2)

    add_para(doc, (
        "We confirm that [IESP Firm Name] and all members of the assessment team are "
        "independent of [Financial Institution Name] and have no financial, business, "
        "or personal relationships that would impair our objectivity or create a conflict "
        "of interest in performing this assessment."
    ))

    add_para(doc, (
        "This engagement has been conducted in accordance with the requirements of "
        "Part C of Appendix 7, BNM RMiT (BNM/RH/PD 028-98, 28 November 2025)."
    ))

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # SECTION 3: DETAIL OF APPLICATION
    # ════════════════════════════════════════════════════════════════════

    add_heading(doc, "Section 3: Detail of Application", level=1)

    # 3.1 Overview
    add_heading(doc, "3.1 Overview of the Application", level=2)

    add_para(doc, (
        "Overview of the application i.e. business case, target segment of "
        "demographic and end-user, etc."
    ), bold=True, size=10, color=DARK_BLUE)

    add_para(doc, (
        "[Provide a concise overview (below 200 words) of the application, including "
        "the business case, target customer segment, end-user demographics, and the "
        "strategic rationale. Additional information may be provided as supporting documents.]"
    ), italic=True)

    add_para(doc, (
        "Describe the technology used to support the product, service or solution."
    ), bold=True, size=10, color=DARK_BLUE)

    add_para(doc, (
        "[Provide a concise description (below 200 words) of the technology stack, "
        "architecture, and key components used to support the service. Include cloud "
        "platforms, frameworks, and integration points. Additional information may be "
        "provided as supporting documents.]"
    ), italic=True)

    # 3.2 Engagement Parameters
    add_heading(doc, "3.2 Engagement Parameters", level=2)

    s3_table = styled_table(doc, rows=1, cols=2, col_widths=[2.8, 3.7])
    s3_table.rows[0].cells[0].text = "Parameter"
    s3_table.rows[0].cells[1].text = "Details"
    style_header_row(s3_table.rows[0])

    s3_fields = [
        ("Engagement type",
         "[Cloud Pre-Implementation / Cloud Attestation / Emerging Tech Pre-Implementation / "
         "Emerging Tech Attestation / DCRA / NRA / Digital Services / BNM Directed]"),
        ("Engagement mode",
         "[Design Adequacy (pre-implementation) / Operating Effectiveness (attestation)]"),
        ("Control sources assessed",
         "[e.g., Appendix 10 (Part A + Part B) + Appendix 7 Part D]"),
        ("BNM trigger clause",
         "[e.g., Paragraph 17.1 / Paragraph 14.1 / Paragraph 16.4]"),
        ("Assessment period",
         "[DD/MM/YYYY] to [DD/MM/YYYY]"),
    ]
    for label, value in s3_fields:
        add_field_row(s3_table, label, value)

    doc.add_paragraph()

    # 3.3 Platforms and Systems in Scope
    add_heading(doc, "3.3 Platforms and Systems in Scope", level=2)

    add_para(doc, "Platforms / CSPs in Scope:", bold=True)

    plat_table = styled_table(doc, rows=4, cols=4, col_widths=[0.5, 2.0, 2.0, 2.0])
    plat_headers = ["#", "Platform / CSP", "Region", "Services Used"]
    for i, h in enumerate(plat_headers):
        plat_table.rows[0].cells[i].text = h
    style_header_row(plat_table.rows[0])

    plat_data = [
        ("1", "[e.g., AWS]", "[e.g., ap-southeast-1]", "[e.g., EC2, S3, RDS, Lambda]"),
        ("2", "[e.g., Microsoft Azure]", "[e.g., Southeast Asia]", "[e.g., AKS, SQL Database]"),
        ("3", "[e.g., On-premise DC]", "[e.g., Cyberjaya]", "[e.g., VMware, NetApp]"),
    ]
    for idx, (num, plat, region, services) in enumerate(plat_data, 1):
        plat_table.rows[idx].cells[0].text = num
        plat_table.rows[idx].cells[1].text = plat
        plat_table.rows[idx].cells[2].text = region
        plat_table.rows[idx].cells[3].text = services
        style_data_row(plat_table.rows[idx], alternate=(idx % 2 == 0))

    doc.add_paragraph()

    add_para(doc, "Critical Systems in Scope:", bold=True)

    sys_table = styled_table(doc, rows=4, cols=5, col_widths=[0.4, 1.8, 1.3, 1.5, 1.5])
    sys_headers = ["#", "System / Application", "Platform", "Criticality", "Data Classification"]
    for i, h in enumerate(sys_headers):
        sys_table.rows[0].cells[i].text = h
    style_header_row(sys_table.rows[0])

    sys_data = [
        ("1", "[e.g., Core Banking]", "[AWS]", "[Critical]", "[Restricted]"),
        ("2", "[e.g., Internet Banking]", "[Azure]", "[Critical]", "[Confidential]"),
        ("3", "[e.g., Mobile App Backend]", "[AWS]", "[High]", "[Confidential]"),
    ]
    for idx, (num, sys, plat, crit, data_cls) in enumerate(sys_data, 1):
        sys_table.rows[idx].cells[0].text = num
        sys_table.rows[idx].cells[1].text = sys
        sys_table.rows[idx].cells[2].text = plat
        sys_table.rows[idx].cells[3].text = crit
        sys_table.rows[idx].cells[4].text = data_cls
        style_data_row(sys_table.rows[idx], alternate=(idx % 2 == 0))

    doc.add_paragraph()

    # 3.4 Scope Exclusions
    add_heading(doc, "3.4 Scope Exclusions", level=2)

    excl_table = styled_table(doc, rows=3, cols=3, col_widths=[0.5, 3.5, 2.5])
    excl_headers = ["#", "Exclusion", "Justification"]
    for i, h in enumerate(excl_headers):
        excl_table.rows[0].cells[i].text = h
    style_header_row(excl_table.rows[0])

    excl_data = [
        ("1", "[e.g., Non-critical development/test environments]",
         "[Not in scope per engagement letter — non-production environments]"),
        ("2", "[e.g., Third-party SaaS applications not hosting critical data]",
         "[Assessed separately under vendor management program]"),
    ]
    for idx, (num, excl, just) in enumerate(excl_data, 1):
        excl_table.rows[idx].cells[0].text = num
        excl_table.rows[idx].cells[1].text = excl
        excl_table.rows[idx].cells[2].text = just
        style_data_row(excl_table.rows[idx], alternate=(idx % 2 == 0))

    doc.add_paragraph()

    # 3.5 Higher-Risk Services Determination
    add_heading(doc, "3.5 Higher-Risk Services Determination", level=2)

    hr_table = styled_table(doc, rows=1, cols=2, col_widths=[3.5, 3.0])
    hr_table.rows[0].cells[0].text = "Criterion"
    hr_table.rows[0].cells[1].text = "Determination"
    style_header_row(hr_table.rows[0])

    hr_fields = [
        ("Customer data processing involved?", "[Yes / No — describe]"),
        ("Cross-border data transmission?", "[Yes / No — describe jurisdictions]"),
        ("Higher-risk services per BNM classification?", "[Yes / No — state basis]"),
        ("Material outsourcing arrangement?", "[Yes / No — describe]"),
    ]
    for label, value in hr_fields:
        add_field_row(hr_table, label, value)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # SECTION 4: TECHNOLOGY RISK ASSESSMENT
    # ════════════════════════════════════════════════════════════════════

    add_heading(doc, "Section 4: Technology Risk Assessment", level=1)

    add_para(doc, (
        "Technology risk assessment shall provide assurance on the effectiveness of "
        "technology risk control and mitigation performed by the financial institution in "
        "meeting expectations outlined in Part D of this Appendix and paragraph 17.1 "
        "(for cloud services and emerging technology)."
    ), italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))

    # 4.1 Executive Summary
    add_heading(doc, "4.1 Executive Summary", level=2)

    # Overall assessment box
    oa_table = styled_table(doc, rows=2, cols=2, col_widths=[2.5, 4.0])
    oa_table.rows[0].cells[0].text = "Overall Assessment"
    oa_table.rows[0].cells[1].text = "Rating"
    style_header_row(oa_table.rows[0])
    oa_table.rows[1].cells[0].text = "Overall Assessment Rating"
    oa_table.rows[1].cells[1].text = "[Satisfactory / Satisfactory with Observations / Requires Improvement / Unsatisfactory]"
    for p in oa_table.rows[1].cells[0].paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
    for p in oa_table.rows[1].cells[1].paragraphs:
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
    set_cell_shading(oa_table.rows[1].cells[1], YELLOW_HEX)

    doc.add_paragraph()

    add_para(doc, (
        "[Paragraph 1: Provide a high-level summary of the assessment scope, including the "
        "engagement type, mode, assessment period, and the control sources assessed. Reference "
        "the applicable BNM RMiT appendix/clauses.]"
    ))

    add_para(doc, (
        "[Paragraph 2: Summarise the key themes from the assessment. Highlight areas of strength "
        "and areas requiring improvement. Reference the number and distribution of findings by "
        "risk rating (High/Medium/Low).]"
    ))

    add_para(doc, (
        "[Paragraph 3: Provide the overall conclusion. State whether the controls assessed are "
        "operating effectively (or designed adequately for pre-implementation) and reference the "
        "Part C attestation type (Type A — Clean / Type B — With Exceptions / Type C — Adverse).]"
    ))

    doc.add_paragraph()

    # 4.2 Assessment Summary by Control Domain
    add_heading(doc, "4.2 Assessment Summary by Control Domain", level=2)

    add_para(doc, (
        "The following table summarises the assessment conclusions by control domain, mapping "
        "directly to the domain rows in the AWP workbook."
    ))

    # Domain summary table
    dom_table = styled_table(doc, rows=1, cols=6,
                             col_widths=[0.8, 1.8, 1.2, 0.8, 1.3, 0.6])
    dom_headers = ["Domain Ref", "Control Domain", "BNM Source", "Level",
                   "Conclusion", "Finding Count"]
    for i, h in enumerate(dom_headers):
        dom_table.rows[0].cells[i].text = h
    style_header_row(dom_table.rows[0])

    # Example domain rows — Cloud engagement
    dom_data = [
        ("CLD-01", "Cloud Risk Management", "App10 Part A", "ORG", "[Compliant]", "[0]"),
        ("CLD-02", "Cloud Strategy & Governance", "App10 Part A", "ORG", "[Compliant]", "[0]"),
        ("CLD-03", "Due Diligence", "App10 Part A", "ORG", "[Partially Compliant]", "[1]"),
        ("CLD-04", "Contractual Arrangements", "App10 Part A", "ORG", "[Compliant]", "[0]"),
        ("CLD-05", "Exit Strategy", "App10 Part A", "ORG", "[Compliant]", "[0]"),
        ("CLD-06", "BNM Consultation", "App10 Part A", "ORG", "[N/A]", "[0]"),
        ("CLD-07", "Board Oversight", "App10 Part A", "ORG", "[Compliant]", "[0]"),
        ("CLD-08", "Cloud Architecture", "App10 Part B", "PLATFORM", "[Partially Compliant]", "[1]"),
        ("CLD-09", "Data Management", "App10 Part B", "PLATFORM", "[Compliant]", "[0]"),
        ("CLD-10", "Access Control", "App10 Part B", "PLATFORM", "[Compliant]", "[0]"),
        ("CLD-11", "Encryption & Key Mgmt", "App10 Part B", "PLATFORM", "[Compliant]", "[0]"),
        ("CLD-12", "Network Security", "App10 Part B", "PLATFORM", "[Non-Compliant]", "[1]"),
        ("...", "...", "...", "...", "...", "..."),
        ("PD-01", "Access Control", "App7 Part D 1(a)", "ORG", "[Compliant]", "[0]"),
        ("PD-02", "Physical & Env Security", "App7 Part D 1(b)", "ORG", "[Compliant]", "[0]"),
        ("...", "...", "...", "...", "...", "..."),
    ]
    for idx, (ref, domain, src, level, conclusion, count) in enumerate(dom_data):
        row = dom_table.add_row()
        row.cells[0].text = ref
        row.cells[1].text = domain
        row.cells[2].text = src
        row.cells[3].text = level
        row.cells[4].text = conclusion
        row.cells[5].text = count
        style_data_row(row, alternate=(idx % 2 == 1))

    add_para(doc, (
        "[Update this table with the actual domain conclusions from the completed AWP workbook. "
        "Each row corresponds to a domain row (bold/green) in the assessment sheet.]"
    ), italic=True, size=9, color=RGBColor(0x99, 0x99, 0x99))

    doc.add_page_break()

    # 4.3 Detailed Findings
    add_heading(doc, "4.3 Detailed Findings", level=2)

    add_para(doc, (
        "This section presents the detailed findings from the assessment. Each finding follows "
        "the condition-criteria-cause-effect-recommendation structure and includes the "
        "management response and agreed remediation actions."
    ))

    # Finding F-001 (High)
    build_finding(
        doc,
        finding_id="F-001",
        title="Insufficient Network Segmentation in Cloud Environment",
        domain="CLD-12 Network Security",
        bnm_ref="App10-B5(a), App10-B5(b)",
        level="PLATFORM",
        platform="AWS ap-southeast-1",
        risk_rating="High",
        conclusion="Non-Compliant",
        condition=(
            "[The FI's AWS environment lacks micro-segmentation between production workloads. "
            "Security groups are configured with overly permissive rules (0.0.0.0/0 ingress on "
            "multiple ports). No network firewall or WAF is deployed for internet-facing workloads. "
            "VPC flow logs are not enabled for all subnets.]"
        ),
        criteria=(
            "[BNM RMiT Appendix 10, Part B, Item 5(a) requires the FI to implement network "
            "segmentation and isolation controls for cloud workloads. Item 5(b) requires "
            "monitoring of network traffic for anomalous activity.]"
        ),
        cause=(
            "[The cloud infrastructure was provisioned using a lift-and-shift approach without "
            "a cloud-native network architecture review. Security group configurations were "
            "carried over from development without hardening for production.]"
        ),
        effect=(
            "[Lateral movement risk in the event of a compromise. Lack of segmentation means "
            "a compromised workload could access other critical systems in the same VPC. "
            "Absence of flow logs impairs incident detection and forensic capability.]"
        ),
        recommendation=(
            "[1. Implement micro-segmentation using security groups with least-privilege rules "
            "within 30 days. 2. Deploy AWS Network Firewall or equivalent for east-west traffic "
            "inspection within 60 days. 3. Enable VPC flow logs for all subnets and integrate "
            "with SIEM within 14 days. 4. Deploy WAF for all internet-facing workloads within "
            "30 days.]"
        ),
        agreed_action="[FI's agreed remediation steps]",
        responsible="[Name, Title]",
        target_date="[DD/MM/YYYY]",
        status="[Open]",
    )

    # Finding F-002 (Medium)
    build_finding(
        doc,
        finding_id="F-002",
        title="Incomplete Cloud Service Provider Due Diligence",
        domain="CLD-03 Due Diligence",
        bnm_ref="App10-A3(a), App10-A3(c)",
        level="ORG",
        platform="All platforms in scope",
        risk_rating="Medium",
        conclusion="Partially Compliant",
        condition=(
            "[The FI's CSP due diligence assessment does not include evaluation of the CSP's "
            "supply chain risk management practices or sub-processor arrangements. The last "
            "comprehensive due diligence review was conducted 18 months ago and does not reflect "
            "current service configurations.]"
        ),
        criteria=(
            "[BNM RMiT Appendix 10, Part A, Item 3(a) requires comprehensive due diligence "
            "of the CSP prior to engagement and periodically thereafter. Item 3(c) requires "
            "assessment of the CSP's ability to comply with Malaysian regulatory requirements.]"
        ),
        cause=(
            "[The due diligence framework was based on a pre-cloud template and has not been "
            "updated to include cloud-specific risk areas such as supply chain, data sovereignty, "
            "and shared responsibility model assessment.]"
        ),
        effect=(
            "[The FI may not have adequate visibility into risks arising from the CSP's "
            "sub-processors or changes to service configurations that could affect data "
            "residency, security controls, or regulatory compliance.]"
        ),
        recommendation=(
            "[1. Update the CSP due diligence framework to include supply chain risk assessment, "
            "sub-processor evaluation, and shared responsibility model review within 60 days. "
            "2. Conduct a refreshed due diligence assessment for all CSPs in scope within 90 days. "
            "3. Establish an annual due diligence refresh cycle.]"
        ),
        agreed_action="[FI's agreed remediation steps]",
        responsible="[Name, Title]",
        target_date="[DD/MM/YYYY]",
        status="[Open]",
    )

    # Finding F-003 (Low)
    build_finding(
        doc,
        finding_id="F-003",
        title="Cloud Architecture Documentation Not Current",
        domain="CLD-08 Cloud Architecture",
        bnm_ref="App10-B1(b)",
        level="PLATFORM",
        platform="AWS ap-southeast-1",
        risk_rating="Low",
        conclusion="Partially Compliant",
        condition=(
            "[The cloud architecture documentation (network diagrams, data flow diagrams, and "
            "system topology) has not been updated since the initial deployment 12 months ago. "
            "Several components added during the period — including a new API gateway and "
            "container orchestration layer — are not reflected in the current documentation.]"
        ),
        criteria=(
            "[BNM RMiT Appendix 10, Part B, Item 1(b) requires the FI to maintain current "
            "and accurate documentation of the cloud architecture, including network topology, "
            "data flows, and system dependencies.]"
        ),
        cause=(
            "[Architecture documentation updates are not included in the change management "
            "process. There is no trigger requiring documentation refresh when infrastructure "
            "changes are deployed.]"
        ),
        effect=(
            "[Outdated architecture documentation may impair incident response, risk assessment, "
            "and audit activities. New components may not have been assessed for security "
            "controls alignment.]"
        ),
        recommendation=(
            "[1. Update all cloud architecture documentation to reflect the current state within "
            "30 days. 2. Integrate documentation update requirements into the change management "
            "process. 3. Establish quarterly architecture documentation review.]"
        ),
        agreed_action="[FI's agreed remediation steps]",
        responsible="[Name, Title]",
        target_date="[DD/MM/YYYY]",
        status="[Open]",
    )

    doc.add_page_break()

    # 4.4 Compliance Summary
    add_heading(doc, "4.4 Compliance Summary", level=2)

    comp_table = styled_table(doc, rows=1, cols=7,
                              col_widths=[1.5, 0.8, 0.8, 0.9, 0.9, 0.5, 0.8])
    comp_headers = ["BNM Source", "Domains Assessed", "Compliant",
                    "Partially Compliant", "Non-Compliant", "N/A", "Compliance %"]
    for i, h in enumerate(comp_headers):
        comp_table.rows[0].cells[i].text = h
    style_header_row(comp_table.rows[0], font_size=8)

    comp_data = [
        ("App10 Part A", "[7]", "[X]", "[X]", "[X]", "[X]", "[X%]"),
        ("App10 Part B", "[14]", "[X]", "[X]", "[X]", "[X]", "[X%]"),
        ("Part D Item 1", "[6]", "[X]", "[X]", "[X]", "[X]", "[X%]"),
        ("Part D Item 2", "[5]", "[X]", "[X]", "[X]", "[X]", "[X%]"),
        ("TOTAL", "[32]", "[X]", "[X]", "[X]", "[X]", "[X%]"),
    ]
    for idx, row_data in enumerate(comp_data):
        row = comp_table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
        style_data_row(row, alternate=(idx % 2 == 1))
        if idx == len(comp_data) - 1:
            # Bold the total row
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True

    add_para(doc, (
        "[Adjust the BNM Source rows based on the actual engagement type. Cloud engagements "
        "use App10 Part A + Part B + Part D. Emerging tech uses App9 + Part D. DCRA uses "
        "Clauses 10.24-10.28 + Part D. NRA uses Clauses 10.36-10.43 + Part D.]"
    ), italic=True, size=9, color=RGBColor(0x99, 0x99, 0x99))

    doc.add_paragraph()

    # 4.5 Findings by Risk Rating
    add_heading(doc, "4.5 Findings by Risk Rating", level=2)

    risk_table = styled_table(doc, rows=1, cols=3, col_widths=[2.0, 2.0, 2.5])
    risk_headers = ["Risk Rating", "Count", "Finding References"]
    for i, h in enumerate(risk_headers):
        risk_table.rows[0].cells[i].text = h
    style_header_row(risk_table.rows[0])

    risk_data = [
        ("High", "[X]", "[F-001, ...]"),
        ("Medium", "[X]", "[F-002, ...]"),
        ("Low", "[X]", "[F-003, ...]"),
        ("TOTAL", "[X]", ""),
    ]
    for idx, (rating, count, refs) in enumerate(risk_data):
        row = risk_table.add_row()
        row.cells[0].text = rating
        row.cells[1].text = count
        row.cells[2].text = refs
        style_data_row(row, alternate=(idx % 2 == 1))
        if idx == len(risk_data) - 1:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # SECTION 5: QUALITY ASSURANCE
    # ════════════════════════════════════════════════════════════════════

    add_heading(doc, "Section 5: Quality Assurance", level=1)

    # Overall recommendation
    add_heading(doc, "5.1 Overall Recommendation", level=2)

    add_para(doc, (
        "[Provide the overall recommendation to the FI and BNM. This should summarise the "
        "assessment outcome, state whether the FI is ready to proceed (for pre-implementation) "
        "or whether controls are operating effectively (for attestation), and highlight any "
        "conditions or prerequisites.]"
    ))

    # Methodology
    add_heading(doc, "5.2 Assessment Methodology", level=2)

    add_para(doc, (
        "The assessment was conducted in accordance with the requirements of Appendix 7, "
        "BNM RMiT (BNM/RH/PD 028-98, 28 November 2025). The methodology comprises:"
    ))

    method_items = [
        "Document review and analysis of policies, procedures, and technical documentation",
        "Interviews with key personnel across technology, security, and risk management functions",
        "Technical validation through system configuration review and evidence inspection",
        "Sample-based testing of controls using judgmental sampling methodology",
        "Re-performance of selected controls to validate operating effectiveness",
        "Direct observation of operational processes and security controls",
    ]
    for item in method_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GREY

    # Sampling
    add_heading(doc, "5.3 Sampling Approach", level=2)

    add_para(doc, (
        "A judgmental (non-statistical) sampling approach was applied, consistent with "
        "limited assurance engagements. Sample sizes were determined based on population "
        "size and control risk classification:"
    ))

    samp_table = styled_table(doc, rows=6, cols=3, col_widths=[2.0, 2.0, 2.5])
    samp_headers = ["Population Size", "Standard Risk", "High Risk"]
    for i, h in enumerate(samp_headers):
        samp_table.rows[0].cells[i].text = h
    style_header_row(samp_table.rows[0])

    samp_data = [
        ("1-5", "All", "All"),
        ("6-15", "3", "5"),
        ("16-50", "5", "8"),
        ("51-100", "8", "12"),
        ("100+", "10", "15"),
    ]
    for idx, (pop, std, high) in enumerate(samp_data, 1):
        samp_table.rows[idx].cells[0].text = pop
        samp_table.rows[idx].cells[1].text = std
        samp_table.rows[idx].cells[2].text = high
        style_data_row(samp_table.rows[idx], alternate=(idx % 2 == 0))

    doc.add_paragraph()

    # Peer review
    add_heading(doc, "5.4 Peer Review", level=2)

    pr_table = styled_table(doc, rows=1, cols=2, col_widths=[2.5, 4.0])
    pr_table.rows[0].cells[0].text = "Item"
    pr_table.rows[0].cells[1].text = "Details"
    style_header_row(pr_table.rows[0])

    pr_fields = [
        ("Peer reviewer name", "[Full name]"),
        ("Qualifications", "[CISA, CISSP, ISO 27001 LA, etc.]"),
        ("Review date", "[DD/MM/YYYY]"),
        ("Review scope", "Methodology, findings, conclusions, and attestation statement"),
        ("Confirmation", "The peer reviewer confirms that the assessment was conducted "
                         "in accordance with professional standards and the findings are "
                         "supported by sufficient and appropriate evidence."),
    ]
    for label, value in pr_fields:
        add_field_row(pr_table, label, value)

    doc.add_paragraph()

    # Scope limitations
    add_heading(doc, "5.5 Scope Limitations and Caveats", level=2)

    add_para(doc, (
        "[Document any scope limitations encountered during the assessment. Examples include: "
        "evidence not provided within the assessment timeline, systems unavailable for testing, "
        "personnel unavailable for interviews, or areas excluded by agreement.]"
    ))

    add_para(doc, (
        "[Document any caveats applicable to the assessment conclusions. This engagement provides "
        "limited assurance based on the procedures performed and evidence gathered during the "
        "assessment period. It does not constitute an audit opinion or absolute assurance.]"
    ))

    # Part C compliance
    add_heading(doc, "5.6 IESP Compliance with Appendix 7 Part C", level=2)

    add_para(doc, (
        "The IESP confirms compliance with all six requirements of Part C "
        "(Requirements on External Party Assurance):"
    ))

    pc_table = styled_table(doc, rows=7, cols=3, col_widths=[0.4, 4.5, 1.6])
    pc_headers = ["#", "Part C Requirement", "Confirmed"]
    for i, h in enumerate(pc_headers):
        pc_table.rows[0].cells[i].text = h
    style_header_row(pc_table.rows[0])

    pc_data = [
        ("1", "Assurance conducted by an independent ESP engaged by the FI",
         "[Yes / No]"),
        ("2", "ESP understands the proposed services, data flows, system architecture, "
              "connectivity, and dependencies",
         "[Yes / No]"),
        ("3", "ESP reviewed comprehensiveness of risk assessment and validated adequacy "
              "of controls",
         "[Yes / No]"),
        ("4", "Report states scope, methodology, findings, and remedial actions (per Part D)",
         "[Yes / No]"),
        ("5", "Report confirms no exception noted based on prescribed risk areas "
              "(negative attestation)",
         "[Yes / No]"),
        ("6", "FI shall provide the report accompanied by relevant documents",
         "[Yes / No]"),
    ]
    for idx, (num, req, confirmed) in enumerate(pc_data, 1):
        pc_table.rows[idx].cells[0].text = num
        pc_table.rows[idx].cells[1].text = req
        pc_table.rows[idx].cells[2].text = confirmed
        style_data_row(pc_table.rows[idx], alternate=(idx % 2 == 0))

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # SECTION 6: AUTHORISED SIGNATORY
    # ════════════════════════════════════════════════════════════════════

    add_heading(doc, "Section 6: Authorised Signatory", level=1)

    add_para(doc, (
        "This report has been prepared by [IESP Firm Name] for the exclusive use of "
        "[Financial Institution Name] and Bank Negara Malaysia. The findings and conclusions "
        "contained herein are based on the procedures performed and evidence gathered during "
        "the assessment period."
    ))

    doc.add_paragraph()
    doc.add_paragraph()

    sig_table = styled_table(doc, rows=4, cols=2, col_widths=[2.0, 4.5])

    sig_data = [
        ("Signature", ""),
        ("Name", "[Full name of authorised signatory]"),
        ("Designation", "[Title / Position]"),
        ("Date", "[DD/MM/YYYY]"),
    ]
    for idx, (label, value) in enumerate(sig_data):
        sig_table.rows[idx].cells[0].text = label
        sig_table.rows[idx].cells[1].text = value
        for p in sig_table.rows[idx].cells[0].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                run.font.color.rgb = DARK_BLUE
        for p in sig_table.rows[idx].cells[1].paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
    # Make signature row taller
    sig_table.rows[0].height = Cm(3)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # APPENDIX I: PART C — NEGATIVE ATTESTATION STATEMENT
    # ════════════════════════════════════════════════════════════════════

    add_heading(doc, "Appendix I: Part C — Negative Attestation Statement", level=1)

    add_para(doc, (
        "The assessor shall select one of the following three attestation types based on "
        "the assessment findings. Delete the types that are not applicable."
    ), italic=True, size=9, color=RGBColor(0x99, 0x99, 0x99))

    doc.add_paragraph()

    # Type A
    add_heading(doc, "Type A: Clean Attestation", level=2)
    set_cell_shading_p = doc.add_paragraph()
    run = set_cell_shading_p.add_run("[Select this type when no material exceptions were identified]")
    run.font.italic = True
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    add_para(doc, (
        "Based on the procedures performed and evidence gathered during the assessment "
        "period [Date] to [Date], nothing has come to our attention that causes us to "
        "believe that the controls assessed are not operating effectively in all material "
        "respects, in accordance with the requirements of [Appendix 10 / Appendix 9 / "
        "Clauses X] and Part D of Appendix 7 of BNM/RH/PD 028-98 (Risk Management in "
        "Technology)."
    ), size=10)

    doc.add_paragraph()

    # Type B
    add_heading(doc, "Type B: Attestation with Exceptions", level=2)
    p = doc.add_paragraph()
    run = p.add_run("[Select this type when material exceptions exist but are not pervasive]")
    run.font.italic = True
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    add_para(doc, (
        "Based on the procedures performed and evidence gathered during the assessment "
        "period [Date] to [Date], except for the matters described in findings "
        "[F-XXX, F-XXX] detailed in Section 4.3 of this report, nothing has come to our "
        "attention that causes us to believe that the controls assessed are not operating "
        "effectively in all material respects, in accordance with the requirements of "
        "[Appendix 10 / Appendix 9 / Clauses X] and Part D of Appendix 7 of BNM/RH/PD "
        "028-98 (Risk Management in Technology)."
    ), size=10)

    doc.add_paragraph()

    # Type C
    add_heading(doc, "Type C: Adverse Attestation", level=2)
    p = doc.add_paragraph()
    run = p.add_run("[Select this type when control failures are significant and pervasive]")
    run.font.italic = True
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    add_para(doc, (
        "Based on the procedures performed and evidence gathered during the assessment "
        "period [Date] to [Date], due to the significance and pervasiveness of the "
        "matters described in Section 4.3 of this report, we are unable to conclude "
        "that the controls assessed are operating effectively in accordance with the "
        "requirements of [Appendix 10 / Appendix 9 / Clauses X] and Part D of "
        "Appendix 7 of BNM/RH/PD 028-98 (Risk Management in Technology)."
    ), size=10)

    doc.add_paragraph()

    # Signature block for attestation
    doc.add_paragraph()

    att_sig = styled_table(doc, rows=4, cols=2, col_widths=[2.0, 4.5])
    att_sig_data = [
        ("Signature", ""),
        ("Name", "[Lead Assessor — full name]"),
        ("Designation", "[Title], [IESP Firm Name]"),
        ("Date", "[DD/MM/YYYY]"),
    ]
    for idx, (label, value) in enumerate(att_sig_data):
        att_sig.rows[idx].cells[0].text = label
        att_sig.rows[idx].cells[1].text = value
        for p in att_sig.rows[idx].cells[0].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                run.font.color.rgb = DARK_BLUE
        for p in att_sig.rows[idx].cells[1].paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
    att_sig.rows[0].height = Cm(3)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # APPENDIX II: PART B — FORMAT OF CONFIRMATION
    # ════════════════════════════════════════════════════════════════════

    add_heading(doc, "Appendix II: Part B — Format of Confirmation", level=1)

    add_para(doc, (
        "This appendix reproduces the verbatim Part B confirmation format from BNM RMiT "
        "Appendix 7 (page 58). This is signed by the FI, not the IESP. The IESP's findings "
        "inform whether the FI can sign this confirmation."
    ), italic=True, size=9, color=RGBColor(0x99, 0x99, 0x99))

    doc.add_paragraph()

    add_para(doc, "Part B: Format of Confirmation", bold=True, size=12, color=DARK_BLUE)

    add_para(doc, "Name of Financial Institution: ............................................................")

    doc.add_paragraph()

    add_para(doc, (
        "As Chairman of the board of directors / designated board-level committee / CISO / "
        "designated senior management officer * of [name of Financial Institution], I confirm "
        "that \u2014"
    ))

    # The 9 points — verbatim from BNM
    confirmations = [
        "digital services / cloud service / emerging technology * is consistent with the "
        "bank\u2019s / insurer\u2019s / takaful operator\u2019s * strategic and business plans;",

        "the board of directors / senior management * understand and are ready to assume "
        "the roles and responsibilities stated in Bank Negara Malaysia\u2019s policy document "
        "on Risk Management in Technology and are also apprised of all relevant provisions "
        "in the FSA, IFSA and DFIA and other relevant legislation, guidelines and codes "
        "of conduct;",

        "risk management process related to digital services / cloud service / emerging "
        "technology * is subject to appropriate oversight by the board of directors and "
        "senior management;",

        "appropriate security measures to address digital services / cloud service / "
        "emerging technology * security concerns are in place;",

        "customer support services and education related to digital services / cloud "
        "service / emerging technology * are in place;",

        "performance monitoring of digital services / cloud service / emerging technology "
        "* products, services, delivery channels and processes has been established;",

        "digital services / cloud service / emerging technology * is included in the "
        "contingency and business resumption plans;",

        "there are adequate resources to support the offering of digital services / cloud "
        "service / emerging technology * business; and",

        "the systems, procedures, security measures, etc. relevant to sound operations "
        "of digital services / cloud service / emerging technology * will be constantly "
        "reviewed to keep up with the latest changes.",
    ]

    for i, text in enumerate(confirmations, 1):
        p = doc.add_paragraph()
        run_num = p.add_run(f"{i}.   ")
        run_num.font.name = 'Calibri'
        run_num.font.size = Pt(10)
        run_num.font.color.rgb = DARK_GREY
        run_text = p.add_run(text)
        run_text.font.name = 'Calibri'
        run_text.font.size = Pt(10)
        run_text.font.color.rgb = DARK_GREY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.3)

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature block
    add_para(doc, "Signature : ............................................")
    add_para(doc, "Name : ............................................")
    add_para(doc, "Date : ............................................")

    doc.add_paragraph()
    add_para(doc, "* (delete whichever is not applicable)", italic=True, size=9)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # APPENDIX III: PART D — MINIMUM CONTROLS SUMMARY
    # ════════════════════════════════════════════════════════════════════

    add_heading(doc, "Appendix III: Part D — Minimum Controls Summary", level=1)

    add_para(doc, (
        "Part D of Appendix 7 defines the minimum controls to be assessed by the IESP. "
        "This appendix summarises the assessment conclusions for all 11 Part D domains."
    ))

    pd_table = styled_table(doc, rows=1, cols=5,
                            col_widths=[0.9, 2.5, 1.0, 1.2, 0.9])
    pd_headers = ["Part D Item", "Description", "Domain Ref", "Conclusion", "Finding Ref"]
    for i, h in enumerate(pd_headers):
        pd_table.rows[0].cells[i].text = h
    style_header_row(pd_table.rows[0])

    pd_data = [
        # Item 1 — Security Requirements
        ("1(a)", "Access control", "PD-01", "[Compliant]", "[N/A]"),
        ("1(b)", "Physical and environmental security", "PD-02", "[Compliant]", "[N/A]"),
        ("1(c)", "Operations security", "PD-03", "[Compliant]", "[N/A]"),
        ("1(d)", "Communication security", "PD-04", "[Compliant]", "[N/A]"),
        ("1(e)", "Information security incident management", "PD-05", "[Compliant]", "[N/A]"),
        ("1(f)", "Information security aspects of business continuity management",
         "PD-06", "[Compliant]", "[N/A]"),
        # Item 2 — Online Transactions and Services
        ("2(a)", "Customer identity authentication", "PD-07", "[Compliant]", "[N/A]"),
        ("2(b)", "Transaction authentication", "PD-08", "[Compliant]", "[N/A]"),
        ("2(c)", "Segregation of duties", "PD-09", "[Compliant]", "[N/A]"),
        ("2(d)", "Data integrity", "PD-10", "[Compliant]", "[N/A]"),
        ("2(e)", "Mobile device risks", "PD-11", "[Compliant]", "[N/A]"),
    ]
    for idx, (item, desc, ref, conclusion, finding) in enumerate(pd_data):
        row = pd_table.add_row()
        row.cells[0].text = item
        row.cells[1].text = desc
        row.cells[2].text = ref
        row.cells[3].text = conclusion
        row.cells[4].text = finding
        style_data_row(row, alternate=(idx % 2 == 1))
        # Highlight Item 1 header area vs Item 2
        if idx == 0:
            # Add section annotation
            pass

    add_para(doc, (
        "[Update conclusions from the completed Appendix 7 Part D sheet in the AWP workbook. "
        "For Item 2 domains, mark as N/A if online transactions and services are not in scope.]"
    ), italic=True, size=9, color=RGBColor(0x99, 0x99, 0x99))

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # APPENDIX IV: DOCUMENT REQUEST LIST
    # ════════════════════════════════════════════════════════════════════

    add_heading(doc, "Appendix IV: Document Request List", level=1)

    add_para(doc, (
        "The following documents and evidence were requested from the financial institution "
        "during the assessment. This list should be issued at engagement commencement and "
        "tracked throughout the assessment period."
    ))

    drl_table = styled_table(doc, rows=1, cols=6,
                             col_widths=[0.3, 2.2, 1.2, 0.5, 0.8, 1.5])
    drl_headers = ["#", "Document / Evidence Requested", "BNM Reference",
                   "Received", "Date Received", "Notes"]
    for i, h in enumerate(drl_headers):
        drl_table.rows[0].cells[i].text = h
    style_header_row(drl_table.rows[0], font_size=8)

    drl_data = [
        ("1", "Board-approved cloud/technology strategy document",
         "App10-A1", "[ ]", "", ""),
        ("2", "Cloud risk management framework / policy",
         "App10-A1", "[ ]", "", ""),
        ("3", "CSP due diligence assessment report(s)",
         "App10-A3", "[ ]", "", ""),
        ("4", "CSP contractual agreements and SLAs",
         "App10-A4", "[ ]", "", ""),
        ("5", "Cloud exit strategy document",
         "App10-A5", "[ ]", "", ""),
        ("6", "Network architecture diagrams (cloud and on-premise)",
         "App10-B5", "[ ]", "", ""),
        ("7", "Data flow diagrams showing cross-border data transmission",
         "App10-B2", "[ ]", "", ""),
        ("8", "IAM policy and privileged access management procedures",
         "App10-B3, App7-D1(a)", "[ ]", "", ""),
        ("9", "Encryption standards and key management procedures",
         "App10-B4", "[ ]", "", ""),
        ("10", "Vulnerability assessment and penetration test reports (last 12 months)",
         "App10-B7", "[ ]", "", ""),
        ("11", "Security incident reports and response logs (last 6 months)",
         "App7-D1(e)", "[ ]", "", ""),
        ("12", "Business continuity plan and DR test results (last 12 months)",
         "App7-D1(f)", "[ ]", "", ""),
        ("13", "Change management policy and recent change logs",
         "App7-D1(c)", "[ ]", "", ""),
        ("14", "Access control matrix and user access review records",
         "App7-D1(a)", "[ ]", "", ""),
        ("15", "Physical security assessment report for data centre(s)",
         "App7-D1(b)", "[ ]", "", ""),
        ("16", "SOC 2 Type II or equivalent report from CSP(s)",
         "App10-A3", "[ ]", "", ""),
        ("17", "Security awareness training records",
         "App7-D1(c)", "[ ]", "", ""),
        ("18", "Third-party risk management framework and vendor register",
         "App8", "[ ]", "", ""),
        ("19", "Board/committee meeting minutes on technology risk (last 4 quarters)",
         "App10-A7", "[ ]", "", ""),
        ("20", "Mobile application security assessment report",
         "App7-D2(e)", "[ ]", "", ""),
    ]
    for idx, (num, doc_name, ref, received, date_r, notes) in enumerate(drl_data):
        row = drl_table.add_row()
        row.cells[0].text = num
        row.cells[1].text = doc_name
        row.cells[2].text = ref
        row.cells[3].text = received
        row.cells[4].text = date_r
        row.cells[5].text = notes
        style_data_row(row, font_size=8, alternate=(idx % 2 == 1))

    doc.add_paragraph()

    add_para(doc, (
        "[Add or remove items based on the specific engagement type and scope. "
        "Cloud engagements typically require items 1-19. DCRA/NRA engagements "
        "will have different document requirements aligned to clauses 10.24-10.28 "
        "or 10.36-10.43 respectively.]"
    ), italic=True, size=9, color=RGBColor(0x99, 0x99, 0x99))

    # ── Save ──
    output_path = "/Users/dawud/claude/iesp/templates/IESP-Report-Template-v2.docx"
    doc.save(output_path)
    print(f"Report template saved to: {output_path}")


if __name__ == "__main__":
    main()
