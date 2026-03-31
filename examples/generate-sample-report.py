#!/usr/bin/env python3
"""
Generate a sample IESP Cloud Pre-Implementation Report as a Word document.

Fictional engagement: Bank Perdana Berhad cloud migration assessment
by CyberAssure Consulting Sdn Bhd.

This is for educational/demonstration purposes only.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colour constants ──
DARK_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
MEDIUM_BLUE = RGBColor(0x2E, 0x75, 0xB6)
LIGHT_BLUE_BG = "D6E4F0"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_RED = RGBColor(0xC0, 0x00, 0x00)
DARK_GREEN = RGBColor(0x00, 0x70, 0x00)
DARK_ORANGE = RGBColor(0xBF, 0x8F, 0x00)
GREY = RGBColor(0x59, 0x56, 0x59)

HIGH_RED = "FFC7CE"
HIGH_RED_TEXT = RGBColor(0x9C, 0x00, 0x06)
MED_YELLOW = "FFEB9C"
MED_YELLOW_TEXT = RGBColor(0x9C, 0x6B, 0x00)
LOW_GREEN = "C6EFCE"
LOW_GREEN_TEXT = RGBColor(0x00, 0x61, 0x00)


def set_cell_shading(cell, color_hex):
    """Set cell background colour."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_formatted_paragraph(doc, text, style="Normal", bold=False, italic=False,
                            size=None, color=None, alignment=None, space_before=None,
                            space_after=None, font_name=None):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
    else:
        run.font.name = "Calibri"
    if alignment is not None:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_section_heading(doc, text, level=1):
    """Add a section heading with dark blue colour."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = DARK_BLUE
        run.font.name = "Calibri"
    return heading


def format_table_header(row, color_hex="1B3A5C"):
    """Format a table header row with dark blue background and white text."""
    for cell in row.cells:
        set_cell_shading(cell, color_hex)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = "Calibri"


def format_table_cell(cell, text, bold=False, size=9, color=None, alignment=None):
    """Format a table cell."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment


def create_simple_table(doc, headers, rows, col_widths=None):
    """Create a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        format_table_cell(cell, header, bold=True)
    format_table_header(table.rows[0])

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            format_table_cell(cell, str(cell_text))
            # Alternate row shading
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def add_finding(doc, ref, severity, title, domain, bnm_ref, condition, criteria,
                cause, effect, recommendation, management_response, target_date):
    """Add a detailed finding."""
    # Finding header
    sev_colors = {"HIGH": DARK_RED, "MEDIUM": DARK_ORANGE, "LOW": DARK_GREEN}
    sev_bg = {"HIGH": HIGH_RED, "MEDIUM": MED_YELLOW, "LOW": LOW_GREEN}

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(f"{ref}: {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"
    p.add_run("  ")
    sev_run = p.add_run(f"  {severity}  ")
    sev_run.bold = True
    sev_run.font.size = Pt(9)
    sev_run.font.color.rgb = sev_colors.get(severity, BLACK)
    sev_run.font.name = "Calibri"

    # Finding details table
    details = [
        ("Domain", domain),
        ("BNM Reference", bnm_ref),
        ("Severity", severity),
    ]
    t = doc.add_table(rows=len(details), cols=2)
    t.style = "Table Grid"
    for i, (label, value) in enumerate(details):
        format_table_cell(t.rows[i].cells[0], label, bold=True, size=9)
        set_cell_shading(t.rows[i].cells[0], LIGHT_BLUE_BG)
        t.rows[i].cells[0].width = Cm(4)
        format_table_cell(t.rows[i].cells[1], value, size=9)

    # Condition
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Condition (What we found):")
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_BLUE
    p2 = doc.add_paragraph(condition)
    for run in p2.runs:
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    # Criteria
    p = doc.add_paragraph()
    run = p.add_run("Criteria (What is required):")
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_BLUE
    p2 = doc.add_paragraph(criteria)
    for run in p2.runs:
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    # Cause
    p = doc.add_paragraph()
    run = p.add_run("Cause:")
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_BLUE
    p2 = doc.add_paragraph(cause)
    for run in p2.runs:
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    # Effect (Risk)
    p = doc.add_paragraph()
    run = p.add_run("Effect (Risk):")
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_BLUE
    p2 = doc.add_paragraph(effect)
    for run in p2.runs:
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    # Recommendation
    p = doc.add_paragraph()
    run = p.add_run("Recommendation:")
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_BLUE
    p2 = doc.add_paragraph(recommendation)
    for run in p2.runs:
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    # Management response table
    resp_table = doc.add_table(rows=3, cols=2)
    resp_table.style = "Table Grid"
    labels = [
        ("Management Response", management_response),
        ("Agreed Action", "Accepted. Remediation plan to be developed."),
        ("Target Completion Date", target_date),
    ]
    for i, (label, value) in enumerate(labels):
        format_table_cell(resp_table.rows[i].cells[0], label, bold=True, size=9)
        set_cell_shading(resp_table.rows[i].cells[0], "E2EFDA")
        resp_table.rows[i].cells[0].width = Cm(4)
        format_table_cell(resp_table.rows[i].cells[1], value, size=9)

    # Separator
    doc.add_paragraph()


def add_footer(section, text="CONFIDENTIAL | IESP-BPB-2026-001 | CyberAssure Consulting Sdn Bhd"):
    """Add footer to a section."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(7.5)
    run.font.color.rgb = GREY
    run.font.name = "Calibri"
    run.italic = True


def build_report():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # Set heading styles
    for i in range(1, 4):
        h_style = doc.styles[f"Heading {i}"]
        h_style.font.name = "Calibri"
        h_style.font.color.rgb = DARK_BLUE

    # ══════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════

    # Add blank space at top
    for _ in range(4):
        doc.add_paragraph()

    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_RED
    run.font.name = "Calibri"

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INDEPENDENT EXTERNAL SERVICE PROVIDER")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TECHNOLOGY RISK ASSESSMENT REPORT")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Cloud Pre-Implementation Review")
    run.font.size = Pt(14)
    run.font.color.rgb = MEDIUM_BLUE
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BNM RMiT Paragraph 17.1 | Appendix 7 Part A Format")
    run.font.size = Pt(11)
    run.font.color.rgb = MEDIUM_BLUE
    run.font.name = "Calibri"

    doc.add_paragraph()
    doc.add_paragraph()

    # Cover details table
    cover_table = doc.add_table(rows=8, cols=2)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_details = [
        ("Financial Institution", "Bank Perdana Berhad"),
        ("IESP Firm", "CyberAssure Consulting Sdn Bhd"),
        ("Engagement Reference", "IESP-BPB-2026-001"),
        ("Engagement Type", "Cloud Pre-Implementation Review (17.1)"),
        ("Assessment Mode", "Design Adequacy"),
        ("Assessment Period", "1 January 2026 to 28 February 2026"),
        ("Report Date", "14 March 2026"),
        ("Report Status", "FINAL"),
    ]
    for i, (label, value) in enumerate(cover_details):
        format_table_cell(cover_table.rows[i].cells[0], label, bold=True, size=10)
        set_cell_shading(cover_table.rows[i].cells[0], LIGHT_BLUE_BG)
        cover_table.rows[i].cells[0].width = Cm(5)
        format_table_cell(cover_table.rows[i].cells[1], value, size=10)
        cover_table.rows[i].cells[1].width = Cm(8)

    # Set borders for cover table
    for row in cover_table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                            top={"val": "single", "sz": "4", "color": "1B3A5C"},
                            bottom={"val": "single", "sz": "4", "color": "1B3A5C"},
                            start={"val": "single", "sz": "4", "color": "1B3A5C"},
                            end={"val": "single", "sz": "4", "color": "1B3A5C"})

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Disclaimer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "This report is confidential and is intended solely for the use of Bank Perdana Berhad "
        "and Bank Negara Malaysia. Unauthorised distribution, copying, or disclosure of any "
        "part of this report is strictly prohibited."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = GREY
    run.font.name = "Calibri"
    run.italic = True

    # Footer
    add_footer(section)

    # ── Page break ──
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # DISCLAIMER PAGE
    # ══════════════════════════════════════════════════════════════

    add_section_heading(doc, "Disclaimer", level=1)

    disclaimer_text = (
        "This report has been prepared by CyberAssure Consulting Sdn Bhd (\"CyberAssure\" or "
        "\"the IESP\") for the exclusive use of Bank Perdana Berhad (\"BPB\" or \"the Bank\") "
        "and Bank Negara Malaysia (\"BNM\") in connection with the Bank's planned migration of "
        "critical systems to cloud infrastructure.\n\n"
        "This report is prepared in accordance with the requirements of the Risk Management in "
        "Technology (RMiT) Policy Document (BNM/RH/PD 028-98, 28 November 2025), specifically "
        "Appendix 7 Part A (Report Format), and is intended to fulfil the pre-implementation "
        "assessment requirement under Paragraph 17.1.\n\n"
        "The assessment was based on information provided by BPB, documentation reviewed, "
        "interviews conducted, and configurations observed during the assessment period. "
        "CyberAssure has not independently verified all information provided and does not "
        "accept liability for any inaccuracies in the information provided by the Bank.\n\n"
        "This report represents a point-in-time assessment. The conclusions and opinions "
        "expressed herein are based on conditions observed during the assessment period and "
        "may not reflect changes occurring after the date of this report.\n\n"
        "The scope of this engagement is limited to design adequacy assessment (pre-implementation). "
        "Operating effectiveness of controls has not been assessed as the systems are not yet "
        "in production."
    )
    p = doc.add_paragraph(disclaimer_text)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════

    add_section_heading(doc, "Table of Contents", level=1)

    toc_items = [
        ("Section 1", "Financial Institution Details"),
        ("Section 2", "External Service Provider Details"),
        ("Section 3", "Scope of Assessment"),
        ("  3.1", "Background and Business Context"),
        ("  3.2", "Technology Description"),
        ("  3.3", "Assessment Scope"),
        ("  3.4", "Assessment Exclusions"),
        ("Section 4", "Technology Risk Assessment"),
        ("  4.1", "Executive Summary"),
        ("  4.2", "Assessment Methodology"),
        ("  4.3", "Detailed Findings"),
        ("  4.4", "Domain Assessment Summary"),
        ("  4.5", "Compliance Summary"),
        ("Section 5", "Quality Assurance"),
        ("  5.1", "Methodology"),
        ("  5.2", "Sampling Approach"),
        ("  5.3", "Peer Review"),
        ("  5.4", "Limitations"),
        ("Section 6", "Authorised Signatory"),
        ("Appendix I", "Part C Negative Attestation Statement"),
        ("Appendix II", "Part B Confirmation Template"),
        ("Appendix III", "Part D Minimum Controls Summary"),
    ]

    for ref, title in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        if ref.startswith("  "):
            run = p.add_run(f"    {ref.strip()}  ")
        else:
            run = p.add_run(f"{ref}  ")
            run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run2 = p.add_run(title)
        run2.font.size = Pt(10)
        run2.font.name = "Calibri"

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECTION 1: FINANCIAL INSTITUTION DETAILS
    # ══════════════════════════════════════════════════════════════

    add_section_heading(doc, "Section 1: Financial Institution Details", level=1)

    fi_data = [
        ("Name of financial institution", "Bank Perdana Berhad"),
        ("Company registration number", "199501018765 (351234-H)"),
        ("Licensed under", "Financial Services Act 2013 (Act 758)"),
        ("Type of licence", "Licensed Bank"),
        ("Registered address",
         "Level 42, Menara Perdana\n"
         "Jalan Sultan Hishamuddin\n"
         "50050 Kuala Lumpur\n"
         "Malaysia"),
        ("Type of application", "Cloud service — New (Pre-Implementation)"),
        ("Description of application",
         "Migration of core banking middleware (CBS-Cloud), "
         "customer data analytics platform (CAP-Cloud), "
         "and API gateway to Amazon Web Services (AWS) "
         "public cloud infrastructure"),
        ("Contact person", "Encik Mohd Faizal bin Ahmad"),
        ("Designation", "Chief Information Security Officer (CISO)"),
        ("Telephone number", "+603-2033-8800"),
        ("Email address", "mohd.faizal@bankperdana.com.my"),
    ]

    t = doc.add_table(rows=len(fi_data), cols=2)
    t.style = "Table Grid"
    for i, (label, value) in enumerate(fi_data):
        format_table_cell(t.rows[i].cells[0], label, bold=True, size=9.5)
        set_cell_shading(t.rows[i].cells[0], LIGHT_BLUE_BG)
        t.rows[i].cells[0].width = Cm(5)
        format_table_cell(t.rows[i].cells[1], value, size=9.5)
        t.rows[i].cells[1].width = Cm(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECTION 2: EXTERNAL SERVICE PROVIDER DETAILS
    # ══════════════════════════════════════════════════════════════

    add_section_heading(doc, "Section 2: External Service Provider Details", level=1)

    iesp_data = [
        ("Name of external service provider", "CyberAssure Consulting Sdn Bhd"),
        ("Company registration number (SSM)", "201401034567 (1098765-U)"),
        ("Registered address",
         "Level 22, Menara CyberPoint\n"
         "Persiaran APEC\n"
         "63000 Cyberjaya, Selangor\n"
         "Malaysia"),
        ("Engagement reference", "IESP-BPB-2026-001"),
        ("Engagement period", "1 January 2026 to 28 February 2026"),
        ("Report date", "14 March 2026"),
        ("Lead assessor", "Ahmad Razif bin Mohd Noor"),
        ("Designation", "Director, Technology Risk Advisory"),
        ("Professional qualifications", "CISA, CISSP, CCSP"),
        ("Telephone number", "+603-8318-2200"),
        ("Email address", "ahmad.razif@cyberassure.com.my"),
    ]

    t = doc.add_table(rows=len(iesp_data), cols=2)
    t.style = "Table Grid"
    for i, (label, value) in enumerate(iesp_data):
        format_table_cell(t.rows[i].cells[0], label, bold=True, size=9.5)
        set_cell_shading(t.rows[i].cells[0], LIGHT_BLUE_BG)
        t.rows[i].cells[0].width = Cm(5)
        format_table_cell(t.rows[i].cells[1], value, size=9.5)
        t.rows[i].cells[1].width = Cm(11)

    # Engagement team
    doc.add_paragraph()
    add_section_heading(doc, "Engagement Team", level=2)

    team_data = [
        ("Ahmad Razif bin Mohd Noor", "Engagement Lead / Director",
         "CISA, CISSP, CCSP", "15 years"),
        ("Nurul Huda binti Ismail", "Senior Consultant",
         "CISA, CCSP, AWS Security Specialty", "8 years"),
        ("Lim Kai Wen", "Consultant",
         "CCSP, AWS Solutions Architect Professional", "5 years"),
        ("Siti Aminah binti Hassan", "Quality Assurance Reviewer",
         "CISA, CISM, CRISC", "18 years"),
    ]

    t = create_simple_table(
        doc,
        ["Name", "Role", "Professional Qualifications", "Experience"],
        team_data,
        col_widths=[5, 4, 4.5, 2.5]
    )

    # Independence confirmation
    doc.add_paragraph()
    add_section_heading(doc, "Independence Confirmation", level=2)

    independence_text = (
        "CyberAssure Consulting Sdn Bhd confirms that the engagement team members are "
        "independent of Bank Perdana Berhad. Specifically:\n\n"
        "(a) No member of the engagement team holds any financial interest in Bank Perdana Berhad "
        "or any of its subsidiaries;\n"
        "(b) No member of the engagement team has been employed by Bank Perdana Berhad within "
        "the preceding 24 months;\n"
        "(c) CyberAssure Consulting has no other active engagements with Bank Perdana Berhad "
        "that could create a conflict of interest;\n"
        "(d) CyberAssure Consulting was not involved in the design, implementation, or selection "
        "of the cloud services under assessment; and\n"
        "(e) The engagement team has adequate understanding of the technology architecture, "
        "cloud service models, and the regulatory requirements of BNM RMiT."
    )
    p = doc.add_paragraph(independence_text)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECTION 3: SCOPE OF ASSESSMENT
    # ══════════════════════════════════════════════════════════════

    add_section_heading(doc, "Section 3: Scope of Assessment", level=1)

    # 3.1 Background
    add_section_heading(doc, "3.1 Background and Business Context", level=2)

    bg_text = (
        "Bank Perdana Berhad (\"BPB\" or \"the Bank\") is a mid-sized Malaysian licensed commercial "
        "bank with total assets of approximately RM42 billion and 1,800 employees across 85 branches "
        "nationwide. As part of its Digital Transformation Strategy 2025-2028 (approved by the Board "
        "of Directors in September 2025), the Bank is undertaking a phased migration of selected "
        "critical systems to public cloud infrastructure.\n\n"
        "The first phase of the migration programme involves three critical systems:\n\n"
        "(i) Core Banking Middleware (CBS-Cloud) — the middleware layer connecting the Bank's "
        "Temenos T24 core banking system to digital channels, partner APIs, and internal applications. "
        "Currently hosted on-premises in the Bank's Cyberjaya data centre.\n\n"
        "(ii) Customer Analytics Platform (CAP-Cloud) — a data analytics and machine learning platform "
        "used for customer segmentation, credit scoring, and campaign management. The platform "
        "processes personally identifiable information (PII) of approximately 2.1 million retail "
        "customers and 45,000 SME customers.\n\n"
        "(iii) API Gateway — the centralised API management platform serving all digital banking "
        "channels (mobile banking, internet banking, open API for fintech partners).\n\n"
        "The Bank has selected Amazon Web Services (AWS) as its primary cloud service provider, "
        "with the ap-southeast-1 (Singapore) region as the primary deployment location. Microsoft "
        "Azure Southeast Asia (Malaysia/Singapore) is designated as the secondary platform for "
        "disaster recovery purposes.\n\n"
        "This migration is classified as a first-time adoption of public cloud for critical systems "
        "under BNM RMiT Paragraph 17.1, triggering the requirement for a pre-implementation IESP "
        "assessment. The Bank is required to consult BNM and submit the IESP report prior to "
        "production migration."
    )
    p = doc.add_paragraph(bg_text)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    # 3.2 Technology Description
    add_section_heading(doc, "3.2 Technology Description", level=2)

    tech_text = (
        "The target architecture deploys the three critical systems on AWS infrastructure with "
        "disaster recovery to Azure:\n\n"
        "Compute and Application Layer: CBS-Cloud runs on Amazon EKS (Kubernetes) with containerised "
        "microservices across three Availability Zones. CAP-Cloud utilises Amazon SageMaker for ML "
        "model training and inference, with Amazon EMR for batch data processing. The API Gateway "
        "is deployed on Amazon API Gateway with AWS Lambda for serverless functions.\n\n"
        "Database Layer: Amazon RDS Aurora PostgreSQL (Multi-AZ) for transactional data, Amazon "
        "Redshift for analytics data warehouse, Amazon DynamoDB for session management and caching.\n\n"
        "Network Connectivity: AWS Direct Connect (2 x 10 Gbps circuits from two carriers) from "
        "Cyberjaya data centre to AWS Singapore. All inter-VPC traffic via AWS Transit Gateway. "
        "VPN backup over public internet with IPSec.\n\n"
        "Security Services: AWS KMS with customer-managed keys (CMK) for encryption at rest. AWS "
        "Shield Advanced for DDoS protection. Amazon GuardDuty for threat detection. AWS CloudTrail "
        "and VPC Flow Logs for audit trail. AWS WAF for web application protection.\n\n"
        "Disaster Recovery: Cross-platform DR to Microsoft Azure Southeast Asia using Azure Site "
        "Recovery and database replication. Target RTO of 4 hours and RPO of 1 hour for CBS-Cloud."
    )
    p = doc.add_paragraph(tech_text)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    # 3.3 Assessment Scope
    add_section_heading(doc, "3.3 Assessment Scope", level=2)

    scope_table_data = [
        ("Regulatory trigger", "BNM RMiT Paragraph 17.1"),
        ("Engagement type", "Cloud Pre-Implementation Review"),
        ("Assessment mode", "Design Adequacy (Pre-Implementation)"),
        ("Control source",
         "Appendix 10 Part A (7 governance areas)\n"
         "Appendix 10 Part B (14 control areas)\n"
         "Appendix 7 Part D (11 minimum control domains)"),
        ("Platforms in scope",
         "AWS ap-southeast-1 (Primary)\n"
         "Azure Southeast Asia (DR)"),
        ("Critical systems in scope",
         "1. Core Banking Middleware (CBS-Cloud)\n"
         "2. Customer Analytics Platform (CAP-Cloud)\n"
         "3. API Gateway"),
        ("Assessment period", "1 January 2026 to 28 February 2026"),
        ("Assessment levels",
         "ORG: Organisation-level controls (assessed once)\n"
         "PLATFORM: Per CSP/platform (AWS + Azure)\n"
         "WORKLOAD: Per critical system (3 systems)"),
    ]

    t = doc.add_table(rows=len(scope_table_data), cols=2)
    t.style = "Table Grid"
    for i, (label, value) in enumerate(scope_table_data):
        format_table_cell(t.rows[i].cells[0], label, bold=True, size=9.5)
        set_cell_shading(t.rows[i].cells[0], LIGHT_BLUE_BG)
        t.rows[i].cells[0].width = Cm(4.5)
        format_table_cell(t.rows[i].cells[1], value, size=9.5)

    # 3.4 Exclusions
    doc.add_paragraph()
    add_section_heading(doc, "3.4 Assessment Exclusions", level=2)

    exclusions = [
        "End-user devices and branch network infrastructure",
        "Third-party integrations beyond the CSP boundary (e.g., RENTAS, PayNet, FPX)",
        "Business process and application logic controls within the T24 core banking system",
        "Financial audit of the cloud migration programme budget",
        "Operating effectiveness testing (this is a design adequacy assessment only — "
        "operating effectiveness will be assessed in the subsequent independent attestation "
        "engagement after production go-live)",
        "Physical security of AWS and Azure data centres (reliance placed on SOC 2 Type II reports)",
    ]
    for item in exclusions:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECTION 4: TECHNOLOGY RISK ASSESSMENT
    # ══════════════════════════════════════════════════════════════

    add_section_heading(doc, "Section 4: Technology Risk Assessment", level=1)

    # 4.1 Executive Summary
    add_section_heading(doc, "4.1 Executive Summary", level=2)

    exec_summary = (
        "CyberAssure Consulting Sdn Bhd has completed its independent pre-implementation "
        "technology risk assessment of Bank Perdana Berhad's planned migration of three critical "
        "systems to AWS and Azure cloud infrastructure."
    )
    p = doc.add_paragraph(exec_summary)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    # Overall Assessment box
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Overall Assessment: Satisfactory with Observations")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Part C Opinion: Type B (With Exceptions)")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = MEDIUM_BLUE
    run.font.name = "Calibri"

    exec_detail = (
        "\nThe Bank has established a generally sound technology risk management framework for "
        "its cloud migration programme. Cloud governance structures are in place, CSP due diligence "
        "has been conducted, and foundational security controls are designed to meet the majority "
        "of Appendix 10 requirements.\n\n"
        "However, the assessment identified two High-severity findings that represent material "
        "exceptions to the Part C negative attestation. These relate to the absence of a "
        "comprehensive cloud exit strategy (F-001) and the lack of a kill-switch mechanism for "
        "the AI/ML analytics platform (F-002). These gaps must be addressed before production "
        "go-live to ensure the Bank's cloud deployment meets the minimum regulatory expectations "
        "under BNM RMiT.\n\n"
        "In addition, four Medium and two Low findings were identified across access controls, "
        "CSP assurance review, infrastructure-as-code security, incident response, skills "
        "management, and backup testing."
    )
    p = doc.add_paragraph(exec_detail)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    # Findings summary table
    summary_table_data = [
        ("HIGH", "2", "F-001, F-002"),
        ("MEDIUM", "4", "F-003, F-004, F-005, F-006"),
        ("LOW", "2", "F-007, F-008"),
        ("Total", "8", ""),
    ]

    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    headers = ["Severity", "Count", "Finding References"]
    for i, h in enumerate(headers):
        format_table_cell(t.rows[0].cells[i], h, bold=True)
    format_table_header(t.rows[0])

    sev_bg_map = {"HIGH": HIGH_RED, "MEDIUM": MED_YELLOW, "LOW": LOW_GREEN, "Total": LIGHT_BLUE_BG}
    for r_idx, (sev, count, refs) in enumerate(summary_table_data):
        format_table_cell(t.rows[r_idx + 1].cells[0], sev, bold=True, size=9)
        format_table_cell(t.rows[r_idx + 1].cells[1], count, size=9,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
        format_table_cell(t.rows[r_idx + 1].cells[2], refs, size=9)
        set_cell_shading(t.rows[r_idx + 1].cells[0], sev_bg_map.get(sev, "FFFFFF"))

    doc.add_paragraph()

    # 4.2 Assessment Methodology
    add_section_heading(doc, "4.2 Assessment Methodology", level=2)

    methodology_text = (
        "The assessment was performed using CyberAssure Consulting's IESP Assessment Methodology "
        "(IAM v3.1), which is aligned with ISACA IT Audit and Assurance Standards, ISO 27001:2022 "
        "audit practices, and the specific requirements of BNM RMiT Appendix 7 Part C.\n\n"
        "The following assessment techniques were employed:\n\n"
        "Documentary Review: Review of 72 documents including policies, architecture designs, "
        "configuration specifications, risk assessments, vendor agreements, and SOC 2 Type II "
        "reports.\n\n"
        "Interviews: 16 interviews conducted with Bank personnel across CISO office, cloud "
        "engineering, application development, risk management, and IT operations.\n\n"
        "Configuration Review: Examination of planned AWS and Azure configurations including "
        "IAM policies, network security groups, encryption settings, logging configurations, "
        "and infrastructure-as-code templates.\n\n"
        "Architecture Walkthrough: Detailed walkthrough of the target architecture with the "
        "cloud engineering team covering compute, database, network, security, and DR components.\n\n"
        "As this is a design adequacy assessment (pre-implementation), the assessment focused on "
        "whether controls are designed to meet regulatory requirements. Operating effectiveness "
        "testing (sampling of operational evidence, re-performance) was not performed and is "
        "not applicable at this stage."
    )
    p = doc.add_paragraph(methodology_text)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    doc.add_page_break()

    # 4.3 Detailed Findings
    add_section_heading(doc, "4.3 Detailed Findings", level=2)

    # F-001
    add_finding(
        doc,
        ref="F-001",
        severity="HIGH",
        title="Incomplete Cloud Exit Strategy",
        domain="CLD-14 Exit Strategy",
        bnm_ref="Appendix 10, Part B, Area 7(a)",
        condition=(
            "The Bank has not developed a comprehensive cloud exit strategy for its planned "
            "AWS deployment. Review of the Cloud Migration Programme documentation and interviews "
            "with the Cloud Programme Manager confirmed that:\n\n"
            "(a) No documented cloud exit strategy exists. The Bank's \"Cloud Exit Considerations\" "
            "document (v0.1, dated October 2025) is a 2-page outline that does not constitute a "
            "strategy.\n"
            "(b) No trigger events have been defined that would activate the exit process (e.g., "
            "CSP insolvency, regulatory direction, material breach of SLA, adverse BNM directive).\n"
            "(c) No migration approach or technical methodology has been defined for extracting "
            "data and workloads from AWS.\n"
            "(d) No timeline, resource requirements, or cost estimates have been prepared for "
            "an exit scenario.\n"
            "(e) Contractual exit provisions are limited to standard AWS Customer Agreement terms "
            "with no bespoke exit support, data retrieval SLAs, or transition assistance clauses."
        ),
        criteria=(
            "Appendix 10 Part B Area 7(a) requires financial institutions to develop an exit "
            "strategy during the planning phase of cloud adoption, not as an afterthought. The "
            "exit strategy must address trigger events, migration approach, timeline, resource "
            "requirements, and contractual provisions for exit support. This requirement exists "
            "to ensure the Bank can transition away from the CSP if circumstances require, "
            "without disruption to critical services."
        ),
        cause=(
            "The cloud programme team prioritised migration design and security controls over "
            "exit planning. The programme plan allocates exit strategy development to Phase 3 "
            "(post-go-live), which does not satisfy the Appendix 10 requirement for exit planning "
            "during the pre-implementation phase."
        ),
        effect=(
            "Without a comprehensive exit strategy, the Bank faces a dependency risk on AWS. In "
            "an adverse scenario (CSP failure, regulatory direction to exit, material contractual "
            "breach), the Bank would not have a tested, documented approach to migrate critical "
            "systems to an alternative provider or back to on-premises within an acceptable "
            "timeframe. This could result in prolonged service disruption, data loss, or "
            "regulatory non-compliance."
        ),
        recommendation=(
            "Develop a comprehensive cloud exit strategy before production go-live, covering at "
            "minimum:\n"
            "(i) Defined trigger events and decision criteria for activating exit;\n"
            "(ii) Data retrieval methodology and technical migration approach;\n"
            "(iii) Target destination (alternative CSP, on-premises, hybrid);\n"
            "(iv) Timeline with milestones and dependencies;\n"
            "(v) Resource requirements (personnel, infrastructure, budget);\n"
            "(vi) Communication plan (internal stakeholders, BNM, customers);\n"
            "(vii) Contractual negotiation with AWS for bespoke exit support and data retrieval SLAs;\n"
            "(viii) Regular testing of exit procedures (annual tabletop exercise as a minimum)."
        ),
        management_response=(
            "Management acknowledges the finding. The cloud programme team will develop a comprehensive "
            "exit strategy as recommended. A dedicated workstream has been established under the "
            "Cloud Programme Director. Draft strategy to be completed by 30 April 2026 and submitted "
            "to BRTC for approval before production go-live."
        ),
        target_date="30 April 2026"
    )

    # F-002
    add_finding(
        doc,
        ref="F-002",
        severity="HIGH",
        title="No Kill-Switch or Suspension Capability for AI Analytics",
        domain="CLD-17 Cybersecurity Operations / AI Governance",
        bnm_ref="Appendix 10, Part B, Area 10(a) / Appendix 9, Item 2(c)",
        condition=(
            "The Customer Analytics Platform (CAP-Cloud) uses machine learning models hosted "
            "on Amazon SageMaker for credit scoring, customer segmentation, and campaign targeting. "
            "The assessment identified that:\n\n"
            "(a) No documented mechanism exists to rapidly suspend or terminate the ML model "
            "inference service if adversarial behaviour, model drift, or biased outcomes are "
            "detected.\n"
            "(b) No defined activation criteria or timeframes exist for triggering a service "
            "suspension (e.g., maximum time from detection to suspension).\n"
            "(c) No fallback process has been documented for manual credit scoring and customer "
            "segmentation in the event the ML service is suspended.\n"
            "(d) The Bank's AI Governance Framework (v1.0, December 2025) references a "
            "\"circuit breaker\" capability but this has not been translated into technical "
            "design or operational procedures.\n"
            "(e) Model monitoring for bias and fairness is planned but not yet configured in "
            "the SageMaker Model Monitor service."
        ),
        criteria=(
            "Appendix 10 Part B Area 10(a) requires adequate cybersecurity controls for cloud "
            "workloads including the ability to respond to and contain security incidents. "
            "Appendix 9 Item 2(c) requires financial institutions adopting AI/emerging technology "
            "to have the ability to suspend or terminate the service when risks materialise. "
            "A kill-switch capability is a fundamental safeguard for AI-driven financial services "
            "where model failures can directly impact customer outcomes and regulatory compliance."
        ),
        cause=(
            "The CAP-Cloud system is being designed by a separate data science team that operates "
            "under a different governance structure from the cloud programme. The kill-switch "
            "requirement was identified in the AI Governance Framework but was not included in "
            "the CAP-Cloud technical requirements specification."
        ),
        effect=(
            "Without a kill-switch capability, a compromised or malfunctioning ML model could "
            "continue making credit decisions or customer segmentation recommendations that are "
            "biased, inaccurate, or manipulated by adversarial inputs. This could result in "
            "unfair treatment of customers, incorrect credit assessments, regulatory sanctions, "
            "and reputational damage."
        ),
        recommendation=(
            "Implement a kill-switch capability for the CAP-Cloud ML service before production "
            "go-live:\n"
            "(i) Design and implement a technical mechanism to suspend ML model inference "
            "within defined timeframes (recommended: < 15 minutes from activation);\n"
            "(ii) Define clear activation criteria covering model drift, bias detection, "
            "adversarial inputs, and anomalous output patterns;\n"
            "(iii) Document and test a fallback process for manual credit scoring and customer "
            "segmentation;\n"
            "(iv) Configure SageMaker Model Monitor for bias, data quality, and model quality "
            "monitoring with automated alerting;\n"
            "(v) Include kill-switch activation in the Bank's incident response playbooks."
        ),
        management_response=(
            "Management accepts the finding. The Data Science team and Cloud Engineering team will "
            "collaborate to design and implement the kill-switch mechanism. SageMaker Model Monitor "
            "configuration will be completed as part of the UAT phase. Fallback procedures will be "
            "documented and tested before go-live."
        ),
        target_date="15 May 2026"
    )

    # F-003
    add_finding(
        doc,
        ref="F-003",
        severity="MEDIUM",
        title="MFA Not Enforced for All Cloud API Access",
        domain="CLD-16 Access Controls",
        bnm_ref="Appendix 10, Part B, Area 9(a)(ii)",
        condition=(
            "Multi-factor authentication (MFA) is enforced for AWS Management Console access via "
            "the Bank's identity provider (Azure AD) with SAML federation. However, the assessment "
            "of the planned IAM configuration identified that:\n\n"
            "(a) Programmatic API access via AWS CLI and SDKs is not subject to MFA enforcement;\n"
            "(b) Twelve (12) service accounts are designed with long-lived access keys (IAM user "
            "credentials) without MFA conditions in the IAM policies;\n"
            "(c) No IAM policy condition (aws:MultiFactorAuthPresent) is applied to restrict "
            "privileged API actions to MFA-authenticated sessions;\n"
            "(d) The planned credential rotation period for service account access keys is 90 days, "
            "which exceeds best practice of using short-lived credentials."
        ),
        criteria=(
            "Appendix 10 Part B Area 9(a)(ii) requires the financial institution to implement "
            "robust access controls for cloud environments including multi-factor authentication "
            "for all privileged access. Industry best practice (CIS AWS Foundations Benchmark v3.0) "
            "recommends eliminating long-lived credentials in favour of IAM roles with temporary "
            "credentials via AWS Security Token Service (STS)."
        ),
        cause=(
            "The cloud engineering team configured MFA for console access but did not extend the "
            "requirement to programmatic access paths. Service accounts were created using the "
            "legacy IAM user model rather than IAM roles with assumed role sessions."
        ),
        effect=(
            "Compromised long-lived access keys could allow unauthorised programmatic access to "
            "cloud resources without MFA challenge. Service account credentials exposed through "
            "code repositories, logs, or workstation compromise could be used to access or "
            "modify cloud infrastructure."
        ),
        recommendation=(
            "Enforce MFA for all cloud access including programmatic API access:\n"
            "(i) Migrate service accounts from IAM users with access keys to IAM roles with "
            "temporary credentials via AWS STS;\n"
            "(ii) Apply MFA condition (aws:MultiFactorAuthPresent) in IAM policies for all "
            "privileged actions;\n"
            "(iii) Implement credential vaulting (e.g., AWS Secrets Manager, HashiCorp Vault) "
            "for any remaining service credentials;\n"
            "(iv) Reduce credential lifespan — adopt short-lived credentials (< 1 hour) where possible."
        ),
        management_response=(
            "Accepted. Cloud engineering team will redesign service account architecture to use "
            "IAM roles and STS. Target completion before production migration."
        ),
        target_date="30 April 2026"
    )

    # F-004
    add_finding(
        doc,
        ref="F-004",
        severity="MEDIUM",
        title="CSP SOC 2 Report Exceptions Not Assessed",
        domain="CLD-04 CSP Certifications and Assurance",
        bnm_ref="Appendix 10, Part A, Area 4(b)",
        condition=(
            "The Bank obtained the AWS SOC 2 Type II report (covering the period 1 October 2024 "
            "to 30 September 2025) as part of its CSP due diligence process. However:\n\n"
            "(a) The Bank has not reviewed or assessed the three (3) control exceptions noted in "
            "the SOC 2 report (relating to access provisioning timeliness, change management "
            "documentation, and physical access review frequency);\n"
            "(b) Complementary User Entity Controls (CUECs) listed in the SOC 2 report have not "
            "been mapped to the Bank's own control environment;\n"
            "(c) No formal assessment has been documented to evaluate the impact of the SOC 2 "
            "exceptions on the Bank's risk posture;\n"
            "(d) The SOC 2 report review was treated as a \"tick-the-box\" exercise — the report "
            "was obtained and filed but not substantively analysed."
        ),
        criteria=(
            "Appendix 10 Part A Area 4(b) requires the financial institution to assess CSP "
            "certifications and third-party assurance reports, including evaluation of any "
            "exceptions or qualifications noted therein. The Bank is expected to understand "
            "the implications of SOC 2 exceptions and implement CUECs to complement the CSP's "
            "control environment."
        ),
        cause=(
            "The risk management team obtained the SOC 2 report but lacked the technical "
            "expertise to assess the exceptions and map CUECs. No formal process exists for "
            "CSP assurance report review."
        ),
        effect=(
            "Unassessed SOC 2 exceptions could indicate control weaknesses at the CSP level "
            "that the Bank has not mitigated through its own controls. Unmapped CUECs mean the "
            "Bank may have gaps in controls that it is responsible for implementing."
        ),
        recommendation=(
            "Establish a formal CSP assurance review process:\n"
            "(i) Review all SOC 2 exceptions and document impact assessment on the Bank's "
            "risk posture;\n"
            "(ii) Map all CUECs to the Bank's control environment and confirm implementation;\n"
            "(iii) Document the review in a CSP Assurance Assessment memorandum;\n"
            "(iv) Establish an annual process for SOC 2 report review and CUEC mapping."
        ),
        management_response=(
            "Accepted. Risk management team will work with the CISO office to conduct a detailed "
            "review of the SOC 2 report exceptions and CUEC mapping. Will be completed before "
            "the BNM consultation submission."
        ),
        target_date="15 April 2026"
    )

    # F-005
    add_finding(
        doc,
        ref="F-005",
        severity="MEDIUM",
        title="IaC Templates Not Scanned for Security Misconfigurations",
        domain="CLD-09 Cloud Application Delivery",
        bnm_ref="Appendix 10, Part B, Area 2(c)(i)",
        condition=(
            "The Bank uses Terraform for infrastructure-as-code (IaC) to manage its AWS cloud "
            "infrastructure. Templates are version-controlled in a GitLab repository with peer "
            "review. However:\n\n"
            "(a) No pre-deployment scanning for security misconfigurations is integrated into the "
            "CI/CD pipeline. Terraform plans are applied after code review but without automated "
            "security validation;\n"
            "(b) Two recent staging deployments (January 2026) resulted in overly permissive "
            "security groups (0.0.0.0/0 ingress on port 443 to internal services) that were "
            "caught during manual review post-deployment;\n"
            "(c) No policy-as-code framework (e.g., Open Policy Agent, Sentinel) is used to "
            "enforce security guardrails on infrastructure configurations;\n"
            "(d) The Terraform state file is stored in S3 with encryption but access controls "
            "to the state bucket have not been reviewed against least-privilege principles."
        ),
        criteria=(
            "Appendix 10 Part B Area 2(c)(i) requires secure configuration management for cloud "
            "deployments. Industry best practice requires IaC scanning as a mandatory gate in "
            "CI/CD pipelines to prevent security misconfigurations from reaching deployed "
            "environments."
        ),
        cause=(
            "The cloud engineering team focused on functional correctness of IaC templates "
            "and did not integrate security scanning tools. The CI/CD pipeline was designed "
            "for speed of deployment rather than security validation."
        ),
        effect=(
            "Without automated IaC scanning, security misconfigurations can be deployed to "
            "cloud environments undetected. Overly permissive security groups, unencrypted "
            "resources, or publicly accessible services could be deployed and remain "
            "undetected until the next manual review."
        ),
        recommendation=(
            "Integrate IaC security scanning into the CI/CD pipeline as a mandatory gate:\n"
            "(i) Deploy an IaC scanning tool (e.g., Checkov, tfsec, Bridgecrew) as a "
            "pipeline stage that must pass before Terraform apply;\n"
            "(ii) Define a baseline security policy covering prohibited configurations "
            "(e.g., no public S3 buckets, no 0.0.0.0/0 security groups, encryption required);\n"
            "(iii) Implement policy-as-code using OPA/Sentinel for custom security guardrails;\n"
            "(iv) Review and restrict access to the Terraform state bucket."
        ),
        management_response=(
            "Accepted. The DevOps team will integrate Checkov into the GitLab CI pipeline and "
            "define baseline security policies. Will be implemented before production deployment."
        ),
        target_date="30 April 2026"
    )

    # F-006
    add_finding(
        doc,
        ref="F-006",
        severity="MEDIUM",
        title="No Cloud-Specific Incident Response Procedures",
        domain="CLD-21 Cyber Response and Recovery",
        bnm_ref="Appendix 10, Part B, Area 14(a)",
        condition=(
            "The Bank has an established Cybersecurity Incident Response Plan (CIRP v4.2, revised "
            "November 2025) that covers its on-premises environment. However, the plan has not "
            "been updated to address cloud-specific scenarios:\n\n"
            "(a) No cloud-specific incident response playbooks exist for scenarios such as AWS "
            "account compromise, cloud data exfiltration, CSP service outage, container escape, "
            "or cross-region failover;\n"
            "(b) The escalation matrix does not include AWS support contacts, AWS Enterprise "
            "Support case creation procedures, or CSP incident coordination;\n"
            "(c) No procedures exist for preserving forensic evidence in cloud environments "
            "(e.g., EBS snapshot before termination, CloudTrail log preservation, VPC Flow Log "
            "retention for investigation);\n"
            "(d) The incident response team has not conducted a tabletop exercise or simulation "
            "for a cloud-specific incident scenario."
        ),
        criteria=(
            "Appendix 10 Part B Area 14(a) requires the financial institution to establish "
            "cyber response and recovery capabilities specific to its cloud deployment. The "
            "Bank's existing IR procedures do not extend to the cloud environment and would "
            "be inadequate for cloud-specific incident scenarios."
        ),
        cause=(
            "The CIRP was last revised before the cloud migration programme commenced. The "
            "incident response team has not yet been trained on cloud-specific incident handling. "
            "Cloud IR playbook development was deprioritised relative to migration activities."
        ),
        effect=(
            "Without cloud-specific IR procedures, the Bank's ability to detect, contain, "
            "investigate, and recover from cloud security incidents is significantly impaired. "
            "Delays in incident response could result in extended data exposure, regulatory "
            "non-compliance with BNM notification timelines, and reputational damage."
        ),
        recommendation=(
            "Develop cloud-specific incident response playbooks before production go-live:\n"
            "(i) Cloud account compromise (IAM credential theft, root account access);\n"
            "(ii) Cloud data exfiltration (S3 data exposure, database breach);\n"
            "(iii) CSP service outage (regional or service-specific);\n"
            "(iv) Container/workload compromise (EKS pod compromise, container escape);\n"
            "(v) Cross-region DR failover activation;\n"
            "(vi) Include AWS Enterprise Support escalation procedures;\n"
            "(vii) Define forensic evidence preservation procedures for cloud;\n"
            "(viii) Conduct a cloud-specific tabletop exercise within 30 days of production go-live."
        ),
        management_response=(
            "Accepted. The CISO office will develop cloud-specific IR playbooks as recommended. "
            "A tabletop exercise will be scheduled within 30 days of production go-live."
        ),
        target_date="31 May 2026"
    )

    # F-007
    add_finding(
        doc,
        ref="F-007",
        severity="LOW",
        title="Cloud Skills Gap Not Formally Assessed",
        domain="CLD-07 Skilled Personnel",
        bnm_ref="Appendix 10, Part A, Area 7(a)(ii)",
        condition=(
            "The cloud operations team comprises four (4) engineers who will support the "
            "production cloud environment. While team members have relevant cloud experience, "
            "the assessment identified that:\n\n"
            "(a) No formal skills gap analysis has been conducted against the competencies "
            "required for operating critical banking systems on AWS;\n"
            "(b) Only one (1) of four (4) cloud engineers holds a CSP-specific certification "
            "(AWS Solutions Architect Associate);\n"
            "(c) No structured training plan with certification targets has been established;\n"
            "(d) The Bank's training budget for cloud skills for FY2026 has not been approved."
        ),
        criteria=(
            "Appendix 10 Part A Area 7(a)(ii) requires the financial institution to ensure "
            "adequate skilled personnel for cloud operations. A formal skills gap analysis "
            "and training plan demonstrates the Bank's commitment to building and maintaining "
            "cloud competency."
        ),
        cause=(
            "The cloud team was assembled from existing IT operations staff with some cloud "
            "experience. A formal competency assessment was not prioritised during the "
            "migration programme."
        ),
        effect=(
            "Without a formal skills assessment and training plan, there is a risk that the "
            "cloud operations team may lack specific competencies required for secure and "
            "efficient operation of critical systems in the cloud. This is a Low-severity "
            "finding as the team has relevant experience but the gap is in formal documentation "
            "and structured development."
        ),
        recommendation=(
            "Conduct a formal skills gap analysis and establish a training plan:\n"
            "(i) Define required competencies for cloud operations (security, networking, "
            "database, containers, monitoring);\n"
            "(ii) Assess each team member against the competency framework;\n"
            "(iii) Establish certification targets (e.g., AWS Solutions Architect, AWS Security "
            "Specialty, CKA for Kubernetes);\n"
            "(iv) Secure FY2026 training budget approval;\n"
            "(v) Track and report progress quarterly to BRTC."
        ),
        management_response=(
            "Accepted. HR and IT will collaborate on the skills gap analysis. Training plan to "
            "be developed and budget to be submitted for FY2026 Q2 approval."
        ),
        target_date="30 June 2026"
    )

    # F-008
    add_finding(
        doc,
        ref="F-008",
        severity="LOW",
        title="Backup Restoration Not Yet Tested for Cloud Workloads",
        domain="CLD-12 Backup and Recovery",
        bnm_ref="Appendix 10, Part B, Area 5(a)(iii)",
        condition=(
            "The Bank's backup strategy for the cloud workloads is defined and technically "
            "configured:\n\n"
            "(a) AWS Backup is configured for automated daily backups of RDS Aurora, EBS "
            "volumes, and DynamoDB tables;\n"
            "(b) Cross-region backup replication to Azure Southeast Asia is configured for "
            "DR purposes;\n"
            "(c) Backup retention policies are defined (daily for 30 days, weekly for 90 days, "
            "monthly for 1 year).\n\n"
            "However, backup restoration testing has not been conducted for any of the three "
            "cloud-hosted systems. No restoration test plan or schedule has been defined."
        ),
        criteria=(
            "Appendix 10 Part B Area 5(a)(iii) requires backup and recovery controls including "
            "regular testing of restoration procedures. While the absence of restoration testing "
            "is understandable at the pre-implementation stage (systems are not yet in production), "
            "restoration testing must be completed before production go-live to validate the "
            "backup strategy."
        ),
        cause=(
            "The backup configuration was completed as part of the infrastructure build. "
            "Restoration testing is planned for the UAT phase but has not yet been scheduled."
        ),
        effect=(
            "Until restoration testing is completed, there is no assurance that backups can be "
            "successfully restored within the defined RTO/RPO targets. This is rated Low as it "
            "is acceptable at the pre-implementation stage but must be addressed before go-live."
        ),
        recommendation=(
            "Conduct backup restoration testing for all three critical systems before production "
            "migration:\n"
            "(i) Test full restoration from AWS Backup for each system;\n"
            "(ii) Test cross-region restoration from Azure DR site;\n"
            "(iii) Validate restored systems against data integrity checks;\n"
            "(iv) Document test results and confirm RTO/RPO targets are achievable;\n"
            "(v) Establish a quarterly restoration test schedule for post-go-live."
        ),
        management_response=(
            "Accepted. Restoration testing will be conducted during the UAT phase (April 2026). "
            "Results will be documented and submitted as part of the go-live readiness assessment."
        ),
        target_date="30 April 2026"
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 4.4 Domain Assessment Summary
    # ══════════════════════════════════════════════════════════════

    add_section_heading(doc, "4.4 Domain Assessment Summary", level=2)

    add_formatted_paragraph(
        doc,
        "The following table summarises the assessment conclusion for each control domain "
        "across Appendix 10 Part A, Part B, and Appendix 7 Part D.",
        size=9.5
    )

    # Part A domains
    add_section_heading(doc, "Appendix 10 Part A: Governance Areas", level=3)

    part_a_domains = [
        ("CLD-01", "Cloud Risk Management", "App10-A1", "Compliant", ""),
        ("CLD-02", "Board and Senior Management Oversight", "App10-A2", "Compliant", ""),
        ("CLD-03", "Cloud Strategy", "App10-A3", "Compliant", ""),
        ("CLD-04", "CSP Certifications and Assurance", "App10-A4", "Partially Compliant", "F-004"),
        ("CLD-05", "Contractual Arrangements", "App10-A5", "Compliant", ""),
        ("CLD-06", "Data Governance", "App10-A6", "Compliant", ""),
        ("CLD-07", "Skilled Personnel", "App10-A7", "Partially Compliant", "F-007"),
    ]

    t = doc.add_table(rows=1 + len(part_a_domains), cols=5)
    t.style = "Table Grid"
    a_headers = ["Ref", "Domain", "BNM Ref", "Conclusion", "Finding"]
    for i, h in enumerate(a_headers):
        format_table_cell(t.rows[0].cells[i], h, bold=True)
    format_table_header(t.rows[0])

    conclusion_colors = {
        "Compliant": LOW_GREEN,
        "Partially Compliant": MED_YELLOW,
        "Non-Compliant": HIGH_RED,
    }

    for r_idx, (ref, domain, bnm, conclusion, finding) in enumerate(part_a_domains):
        row = t.rows[r_idx + 1]
        format_table_cell(row.cells[0], ref, size=8.5)
        format_table_cell(row.cells[1], domain, size=8.5)
        format_table_cell(row.cells[2], bnm, size=8.5)
        format_table_cell(row.cells[3], conclusion, size=8.5, bold=True)
        if conclusion in conclusion_colors:
            set_cell_shading(row.cells[3], conclusion_colors[conclusion])
        format_table_cell(row.cells[4], finding, size=8.5)

    doc.add_paragraph()

    # Part B domains
    add_section_heading(doc, "Appendix 10 Part B: Control Areas", level=3)

    part_b_domains = [
        ("CLD-08", "Cloud Architecture and Design", "App10-B1", "Compliant", ""),
        ("CLD-09", "Cloud Application Delivery", "App10-B2", "Partially Compliant", "F-005"),
        ("CLD-10", "Cloud Data Protection", "App10-B3", "Compliant", ""),
        ("CLD-11", "Cloud Network Security", "App10-B4", "Compliant", ""),
        ("CLD-12", "Backup and Recovery", "App10-B5", "Partially Compliant", "F-008"),
        ("CLD-13", "Cloud Monitoring and Logging", "App10-B6", "Compliant", ""),
        ("CLD-14", "Exit Strategy", "App10-B7", "Non-Compliant", "F-001"),
        ("CLD-15", "Business Continuity Management", "App10-B8", "Compliant", ""),
        ("CLD-16", "Access Controls", "App10-B9", "Partially Compliant", "F-003"),
        ("CLD-17", "Cybersecurity Operations", "App10-B10", "Non-Compliant", "F-002"),
        ("CLD-18", "Vulnerability Management", "App10-B11", "Compliant", ""),
        ("CLD-19", "Change Management", "App10-B12", "Compliant", ""),
        ("CLD-20", "CSP Service Management", "App10-B13", "Compliant", ""),
        ("CLD-21", "Cyber Response and Recovery", "App10-B14", "Partially Compliant", "F-006"),
    ]

    t = doc.add_table(rows=1 + len(part_b_domains), cols=5)
    t.style = "Table Grid"
    for i, h in enumerate(a_headers):
        format_table_cell(t.rows[0].cells[i], h, bold=True)
    format_table_header(t.rows[0])

    for r_idx, (ref, domain, bnm, conclusion, finding) in enumerate(part_b_domains):
        row = t.rows[r_idx + 1]
        format_table_cell(row.cells[0], ref, size=8.5)
        format_table_cell(row.cells[1], domain, size=8.5)
        format_table_cell(row.cells[2], bnm, size=8.5)
        format_table_cell(row.cells[3], conclusion, size=8.5, bold=True)
        if conclusion in conclusion_colors:
            set_cell_shading(row.cells[3], conclusion_colors[conclusion])
        format_table_cell(row.cells[4], finding, size=8.5)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 4.5 Compliance Summary
    # ══════════════════════════════════════════════════════════════

    add_section_heading(doc, "4.5 Compliance Summary", level=2)

    compliance_data = [
        ("Appendix 10 Part A (Governance)", "7", "5", "2", "0", "71%"),
        ("Appendix 10 Part B (Controls)", "14", "8", "4", "2", "57%"),
        ("Appendix 7 Part D (Minimum Controls)", "11", "10", "1", "0", "91%"),
        ("Total", "32", "23", "7", "2", "72%"),
    ]

    t = doc.add_table(rows=1 + len(compliance_data), cols=6)
    t.style = "Table Grid"
    comp_headers = ["Assessment Area", "Total Domains", "Compliant",
                    "Partially Compliant", "Non-Compliant", "Compliance %"]
    for i, h in enumerate(comp_headers):
        format_table_cell(t.rows[0].cells[i], h, bold=True)
    format_table_header(t.rows[0])

    for r_idx, row_data in enumerate(compliance_data):
        row = t.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            bold = (r_idx == len(compliance_data) - 1)
            format_table_cell(row.cells[c_idx], val, size=9, bold=bold,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else None)
            if r_idx == len(compliance_data) - 1:
                set_cell_shading(row.cells[c_idx], LIGHT_BLUE_BG)

    doc.add_paragraph()

    # Compliance commentary
    commentary = (
        "The overall compliance rate of 72% reflects a Satisfactory with Observations position. "
        "The two Non-Compliant domains (Exit Strategy and Cybersecurity Operations/AI Governance) "
        "are the basis for the Type B (With Exceptions) Part C opinion. The seven Partially "
        "Compliant domains have gaps that are addressable before production go-live.\n\n"
        "Part A Governance (6 of 7 Compliant): The governance framework for cloud adoption is "
        "generally well-established. The Board-approved cloud strategy, risk management framework, "
        "and contractual arrangements demonstrate appropriate governance oversight. The two partial "
        "compliance areas relate to CSP assurance review depth (F-004) and cloud skills formalisation "
        "(F-007).\n\n"
        "Part B Controls (8 of 14 Compliant): Core technical controls for cloud architecture, "
        "data protection, network security, and monitoring are adequately designed. The material "
        "gaps in exit strategy (F-001) and AI kill-switch capability (F-002) are the primary "
        "concerns. Additional gaps in access controls (F-003), IaC security (F-005), incident "
        "response (F-006), and backup testing (F-008) are Medium/Low severity and addressable.\n\n"
        "Part D Minimum Controls (10 of 11 Compliant): The minimum controls baseline is largely "
        "met. The one Partially Compliant domain relates to information security incident management "
        "(linked to the cloud-specific IR gap in F-006)."
    )
    p = doc.add_paragraph(commentary)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECTION 5: QUALITY ASSURANCE
    # ══════════════════════════════════════════════════════════════

    add_section_heading(doc, "Section 5: Quality Assurance", level=1)

    # 5.1 Methodology
    add_section_heading(doc, "5.1 Methodology", level=2)

    meth_text = (
        "This assessment was performed in accordance with the following standards and frameworks:\n\n"
        "(a) BNM RMiT (BNM/RH/PD 028-98, 28 November 2025), specifically Appendix 7 Parts A, C, "
        "and D, and Appendix 10 Parts A and B;\n"
        "(b) ISACA IT Audit and Assurance Standards (ITAAS);\n"
        "(c) ISO 27001:2022 audit practices;\n"
        "(d) CyberAssure Consulting's IESP Assessment Methodology (IAM v3.1).\n\n"
        "The assessment methodology addresses the six requirements of Appendix 7 Part C:\n\n"
        "Part C Requirement 1 (Independence): CyberAssure is an independent external service "
        "provider with no financial interest in, employment relationship with, or advisory role "
        "in the design or implementation of the systems under assessment. Independence confirmed "
        "in Section 2.\n\n"
        "Part C Requirement 2 (Understanding): The engagement team demonstrated understanding of "
        "the proposed cloud services, data flows, system architecture, connectivity, and "
        "dependencies through detailed architecture walkthroughs, documentation review, and "
        "interviews.\n\n"
        "Part C Requirement 3 (Comprehensiveness): The risk assessment covers all areas specified "
        "in Appendix 10 (21 governance and control areas) and Appendix 7 Part D (11 minimum "
        "control domains).\n\n"
        "Part C Requirement 4 (Reporting): This report states the scope, methodology, findings, "
        "and remedial actions as required.\n\n"
        "Part C Requirement 5 (Attestation): The Part C negative attestation is provided in "
        "Appendix I.\n\n"
        "Part C Requirement 6 (Documentation): All supporting documentation, working papers, and "
        "evidence are maintained and available for BNM review upon request."
    )
    p = doc.add_paragraph(meth_text)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    # 5.2 Sampling
    add_section_heading(doc, "5.2 Sampling Approach", level=2)

    sampling_text = (
        "Judgmental (non-statistical) sampling was employed, which is appropriate for limited "
        "assurance engagements under ISACA standards. As this is a design adequacy assessment "
        "(pre-implementation), sampling was applied to documentary evidence rather than "
        "operational transactions.\n\n"
        "Documents reviewed: 72 (policies, architecture documents, configuration specifications, "
        "risk assessments, vendor agreements, SOC 2 reports)\n"
        "Interviews conducted: 16 (with Bank personnel across 6 functions)\n"
        "Configuration reviews: Planned configurations for 3 critical systems across 2 CSPs\n"
        "Architecture walkthroughs: 4 sessions covering compute, database, network/security, "
        "and DR architecture"
    )
    p = doc.add_paragraph(sampling_text)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    # 5.3 Peer Review
    add_section_heading(doc, "5.3 Peer Review", level=2)

    peer_text = (
        "In accordance with CyberAssure Consulting's quality assurance procedures and the "
        "requirements of BNM RMiT Appendix 7 Part C, this report has been subjected to "
        "independent peer review.\n\n"
        "Peer Reviewer: Puan Siti Aminah binti Hassan\n"
        "Designation: Partner, Technology Risk Advisory\n"
        "Professional Qualifications: CISA, CISM, CRISC\n"
        "Experience: 18 years in technology risk advisory and IESP assessments\n\n"
        "The peer review covered:\n"
        "(a) Adequacy of scope relative to regulatory requirements;\n"
        "(b) Sufficiency and appropriateness of evidence gathered;\n"
        "(c) Reasonableness of conclusions and findings;\n"
        "(d) Compliance with Appendix 7 Part A reporting format;\n"
        "(e) Accuracy and completeness of the Part C negative attestation.\n\n"
        "The peer review was completed on 10 March 2026. All review comments have been addressed "
        "and incorporated into this final report."
    )
    p = doc.add_paragraph(peer_text)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    # 5.4 Limitations
    add_section_heading(doc, "5.4 Limitations", level=2)

    limitations_text = (
        "The following limitations apply to this assessment:\n\n"
        "(a) Design adequacy only: This assessment evaluates whether controls are designed to meet "
        "regulatory requirements. Operating effectiveness of controls has not been assessed as the "
        "systems are not yet in production. A separate independent attestation engagement will be "
        "required to assess operating effectiveness after production go-live.\n\n"
        "(b) Point-in-time assessment: The conclusions are based on conditions observed during the "
        "assessment period (1 January to 28 February 2026). Changes to the architecture, "
        "configurations, or control environment after this date are not reflected.\n\n"
        "(c) Information reliance: The assessment relied on information provided by the Bank. "
        "While the engagement team exercised professional scepticism, CyberAssure has not "
        "independently verified all information provided.\n\n"
        "(d) CSP physical controls: Physical security of AWS and Azure data centres was not "
        "directly assessed. Reliance was placed on SOC 2 Type II reports for CSP physical "
        "and environmental controls.\n\n"
        "(e) Limited assurance: This engagement provides limited assurance (negative attestation) "
        "and does not constitute an audit opinion or reasonable assurance engagement."
    )
    p = doc.add_paragraph(limitations_text)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SECTION 6: AUTHORISED SIGNATORY
    # ══════════════════════════════════════════════════════════════

    add_section_heading(doc, "Section 6: Authorised Signatory", level=1)

    signatory_text = (
        "This report has been prepared and is issued under the authority of the undersigned."
    )
    p = doc.add_paragraph(signatory_text)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Signature block
    p = doc.add_paragraph()
    run = p.add_run("_" * 40)
    run.font.name = "Calibri"

    sig_details = [
        ("Name:", "Ahmad Razif bin Mohd Noor"),
        ("Designation:", "Director, Technology Risk Advisory"),
        ("Firm:", "CyberAssure Consulting Sdn Bhd"),
        ("Professional Qualifications:", "CISA, CISSP, CCSP"),
        ("Date:", "14 March 2026"),
    ]

    for label, value in sig_details:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(label + " ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run2 = p.add_run(value)
        run2.font.size = Pt(10)
        run2.font.name = "Calibri"

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # APPENDIX I: PART C NEGATIVE ATTESTATION
    # ══════════════════════════════════════════════════════════════

    add_section_heading(doc, "Appendix I: Part C Negative Attestation Statement", level=1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("TYPE B ATTESTATION — WITH EXCEPTIONS")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"

    doc.add_paragraph()

    attestation_text = (
        "To: The Board of Directors, Bank Perdana Berhad\n"
        "     Bank Negara Malaysia\n\n"
        "Independent External Service Provider Negative Attestation Statement\n\n"
        "Pursuant to the Risk Management in Technology (RMiT) Policy Document "
        "(BNM/RH/PD 028-98, 28 November 2025), Appendix 7 Part C, CyberAssure Consulting "
        "Sdn Bhd, as the appointed Independent External Service Provider, hereby provides "
        "the following negative attestation in respect of the technology risk controls "
        "assessment for Bank Perdana Berhad's planned migration of critical systems to "
        "cloud infrastructure.\n\n"
        "Engagement Reference: IESP-BPB-2026-001\n"
        "Assessment Period: 1 January 2026 to 28 February 2026\n"
        "Assessment Mode: Design Adequacy (Pre-Implementation)\n"
        "Systems Assessed: Core Banking Middleware (CBS-Cloud), Customer Analytics Platform "
        "(CAP-Cloud), API Gateway\n"
        "Cloud Platforms: AWS ap-southeast-1 (Primary), Azure Southeast Asia (DR)\n\n"
    )
    p = doc.add_paragraph(attestation_text)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = "Calibri"

    # The attestation statement
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Attestation Statement:")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"

    statement = (
        "\nBased on the procedures performed and the evidence obtained during the assessment "
        "period, except for the matters described below, nothing has come to our attention "
        "that causes us to believe that the technology risk management controls assessed are "
        "not designed to operate effectively in accordance with the requirements of BNM RMiT "
        "Appendix 10 and Appendix 7 Part D.\n"
    )
    p = doc.add_paragraph(statement)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = "Calibri"

    # Exceptions
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Material Exceptions:")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_RED
    run.font.name = "Calibri"

    exceptions = [
        (
            "F-001 — Incomplete Cloud Exit Strategy: The Bank has not developed a comprehensive "
            "cloud exit strategy as required by Appendix 10 Part B Area 7(a). No documented "
            "trigger events, migration approach, timeline, or resource requirements exist for "
            "an exit scenario. This represents a material gap in the Bank's preparedness to "
            "transition away from the CSP if required."
        ),
        (
            "F-002 — No Kill-Switch or Suspension Capability for AI Analytics: The Customer "
            "Analytics Platform, which uses ML models for credit scoring and customer "
            "segmentation, has no documented mechanism to rapidly suspend the service if "
            "adversarial behaviour, model drift, or biased outcomes are detected, as required "
            "by Appendix 10 Part B Area 10(a) and Appendix 9 Item 2(c). This represents a "
            "material gap in the Bank's ability to control AI-driven financial services."
        ),
    ]

    for i, exc in enumerate(exceptions, 1):
        p = doc.add_paragraph(f"{i}. {exc}")
        for run in p.runs:
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"

    doc.add_paragraph()

    impact_text = (
        "These exceptions are material but not pervasive. The remaining 30 of 32 assessed "
        "control domains are Compliant or Partially Compliant, with the Partially Compliant "
        "domains having gaps that are addressable before production go-live. Accordingly, "
        "a Type B (With Exceptions) opinion is issued rather than Type C (Adverse).\n\n"
        "We recommend that the Bank remediates both material exceptions before production "
        "go-live and prior to seeking BNM approval for the cloud deployment."
    )
    p = doc.add_paragraph(impact_text)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature
    p = doc.add_paragraph()
    run = p.add_run("_" * 40)
    run.font.name = "Calibri"

    sig_attestation = [
        ("Name:", "Ahmad Razif bin Mohd Noor"),
        ("Designation:", "Director, Technology Risk Advisory"),
        ("Firm:", "CyberAssure Consulting Sdn Bhd"),
        ("Date:", "14 March 2026"),
    ]

    for label, value in sig_attestation:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(label + " ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run2 = p.add_run(value)
        run2.font.size = Pt(10)
        run2.font.name = "Calibri"

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # APPENDIX II: PART B CONFIRMATION TEMPLATE
    # ══════════════════════════════════════════════════════════════

    add_section_heading(doc, "Appendix II: Part B Confirmation Template", level=1)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        "The following is the Appendix 7 Part B confirmation template for Bank Perdana Berhad "
        "to complete after remediation of the findings in this report. This confirmation is "
        "signed by the Bank's CISO or designated board committee chair and submitted to BNM "
        "as part of the cloud consultation package."
    )
    run.font.size = Pt(9.5)
    run.font.name = "Calibri"
    run.italic = True

    doc.add_paragraph()

    # Part B header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FORMAT OF CONFIRMATION\n(APPENDIX 7 PART B)")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"

    doc.add_paragraph()

    part_b_intro = (
        "To: Bank Negara Malaysia\n\n"
        "Re: Cloud Pre-Implementation — Bank Perdana Berhad\n"
        "Engagement Reference: IESP-BPB-2026-001\n\n"
        "We, the undersigned, hereby confirm that in respect of the planned migration of "
        "critical systems (Core Banking Middleware, Customer Analytics Platform, API Gateway) "
        "to cloud infrastructure (AWS and Azure):"
    )
    p = doc.add_paragraph(part_b_intro)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = "Calibri"

    # 9 attestation items
    part_b_items = [
        "The cloud adoption is consistent with the Bank's strategic and business plans;",
        "The Board of Directors and senior management understand and accept the responsibilities "
        "under the Risk Management in Technology (RMiT) Policy Document;",
        "The technology risk management process is subject to appropriate Board oversight;",
        "Appropriate security measures are in place to protect customer information and the "
        "Bank's technology assets in the cloud environment;",
        "Adequate customer support services are in place for services delivered through the "
        "cloud platform;",
        "Performance monitoring mechanisms have been established for the cloud services;",
        "The cloud systems have been included in the Bank's contingency and business resumption "
        "plans;",
        "Adequate resources (personnel, budget, infrastructure) are available to support the "
        "cloud deployment and ongoing operations;",
        "Systems, procedures, and security measures will be regularly reviewed and updated to "
        "keep current with the evolving threat landscape and regulatory requirements.",
    ]

    for i, item in enumerate(part_b_items, 1):
        p = doc.add_paragraph(f"{i}. {item}")
        for run in p.runs:
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"

    doc.add_paragraph()
    doc.add_paragraph()

    # Blank signature block
    p = doc.add_paragraph()
    run = p.add_run("_" * 40)
    run.font.name = "Calibri"

    blank_sig = [
        ("Name:", "[To be completed by Bank Perdana Berhad]"),
        ("Designation:", "[CISO / Board Risk & Technology Committee Chair]"),
        ("Date:", "[To be completed after remediation]"),
    ]

    for label, value in blank_sig:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(label + " ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run2 = p.add_run(value)
        run2.font.size = Pt(10)
        run2.font.name = "Calibri"
        run2.italic = True
        run2.font.color.rgb = GREY

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # APPENDIX III: PART D MINIMUM CONTROLS SUMMARY
    # ══════════════════════════════════════════════════════════════

    add_section_heading(doc, "Appendix III: Appendix 7 Part D — Minimum Controls Summary", level=1)

    add_formatted_paragraph(
        doc,
        "The following table summarises the assessment of the 11 minimum control domains "
        "specified in Appendix 7 Part D, which applies to all IESP engagements as a baseline.",
        size=9.5
    )

    doc.add_paragraph()

    # Part D Item 1: Security Requirements
    add_section_heading(doc, "Item 1: Security Requirements", level=2)

    part_d_item1 = [
        ("PD-01", "1(a)", "Access Control", "Compliant",
         "Robust access control framework designed. IAM policies, RBAC, and privileged access "
         "management controls are in place. MFA gap noted in F-003 but assessed under Appendix 10."),
        ("PD-02", "1(b)", "Physical and Environmental Security", "Compliant",
         "Reliance placed on AWS SOC 2 Type II report and Azure SOC 2 report for CSP physical "
         "security. Bank's Cyberjaya data centre (Direct Connect termination) has adequate "
         "physical controls."),
        ("PD-03", "1(c)", "Operations Security", "Compliant",
         "Cloud operations security controls are designed including change management, "
         "capacity management, and monitoring. IaC and operational procedures documented."),
        ("PD-04", "1(d)", "Communication Security", "Compliant",
         "Network security controls designed with encryption in transit (TLS 1.2+), Direct "
         "Connect with MACsec, VPN backup, and Transit Gateway for inter-VPC traffic."),
        ("PD-05", "1(e)", "Information Security Incident Management", "Partially Compliant",
         "Incident management framework exists but cloud-specific playbooks are not yet "
         "developed (linked to F-006). IR capability for on-premises is adequate."),
        ("PD-06", "1(f)", "Information Security Aspects of BCM", "Compliant",
         "Cloud DR architecture designed with cross-region/cross-platform redundancy. "
         "BCP updated to include cloud scenarios. Restoration testing pending (F-008)."),
    ]

    t = doc.add_table(rows=1 + len(part_d_item1), cols=5)
    t.style = "Table Grid"
    pd_headers = ["Ref", "Part D", "Domain", "Conclusion", "Key Observations"]
    for i, h in enumerate(pd_headers):
        format_table_cell(t.rows[0].cells[i], h, bold=True)
    format_table_header(t.rows[0])

    for r_idx, (ref, pd_ref, domain, conclusion, obs) in enumerate(part_d_item1):
        row = t.rows[r_idx + 1]
        format_table_cell(row.cells[0], ref, size=8)
        format_table_cell(row.cells[1], pd_ref, size=8)
        format_table_cell(row.cells[2], domain, size=8)
        format_table_cell(row.cells[3], conclusion, size=8, bold=True)
        if conclusion in conclusion_colors:
            set_cell_shading(row.cells[3], conclusion_colors[conclusion])
        format_table_cell(row.cells[4], obs, size=8)

    doc.add_paragraph()

    # Part D Item 2: Online Transactions and Services
    add_section_heading(doc, "Item 2: Online Transactions and Services", level=2)

    part_d_item2 = [
        ("PD-07", "2(a)", "Customer Identity Authentication", "Compliant",
         "Multi-factor authentication for internet and mobile banking. Biometric and OTP "
         "authentication designed for cloud-hosted channels."),
        ("PD-08", "2(b)", "Transaction Authentication", "Compliant",
         "Transaction signing with OTP and push notification. Risk-based authentication "
         "for high-value transactions designed."),
        ("PD-09", "2(c)", "Segregation of Duties", "Compliant",
         "Role-based access control with maker-checker for administrative functions. "
         "Segregation between development, staging, and production environments."),
        ("PD-10", "2(d)", "Data Integrity", "Compliant",
         "Data integrity controls designed including database checksums, API input "
         "validation, and audit trail for data modifications."),
        ("PD-11", "2(e)", "Mobile Device Risks", "Compliant",
         "Mobile application security controls designed including certificate pinning, "
         "jailbreak/root detection, secure key storage, and runtime protection."),
    ]

    t = doc.add_table(rows=1 + len(part_d_item2), cols=5)
    t.style = "Table Grid"
    for i, h in enumerate(pd_headers):
        format_table_cell(t.rows[0].cells[i], h, bold=True)
    format_table_header(t.rows[0])

    for r_idx, (ref, pd_ref, domain, conclusion, obs) in enumerate(part_d_item2):
        row = t.rows[r_idx + 1]
        format_table_cell(row.cells[0], ref, size=8)
        format_table_cell(row.cells[1], pd_ref, size=8)
        format_table_cell(row.cells[2], domain, size=8)
        format_table_cell(row.cells[3], conclusion, size=8, bold=True)
        if conclusion in conclusion_colors:
            set_cell_shading(row.cells[3], conclusion_colors[conclusion])
        format_table_cell(row.cells[4], obs, size=8)

    doc.add_paragraph()

    # Part D summary
    pd_summary = (
        "Part D Minimum Controls Summary: 10 of 11 domains are Compliant. One domain "
        "(PD-05: Information Security Incident Management) is Partially Compliant due to "
        "the absence of cloud-specific incident response playbooks (linked to Finding F-006). "
        "The overall Part D compliance rate of 91% indicates a strong baseline control posture."
    )
    p = doc.add_paragraph(pd_summary)
    for run in p.runs:
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"
        run.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # End of report
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— END OF REPORT —")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"

    # ── Save ──
    output_path = "/Users/dawud/claude/iesp/examples/Sample-IESP-Cloud-Report.docx"
    doc.save(output_path)
    print(f"Report saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    build_report()
